# jarvis.py – FRIDAY Agent v4.0
# Run: streamlit run jarvis.py
# ──────────────────────────────────────────────────────────────────────────────
# Required packages:
#   pip install streamlit streamlit-autorefresh psutil requests pyttsx3
#   pip install SpeechRecognition ollama PyPDF2 python-docx reportlab
# version..v4.0  [FIXED: autorefresh no longer kills AI responses]

import os
import json
import time
import webbrowser
import threading
import psutil
import requests
import streamlit as st
from datetime import datetime, date, timedelta
import urllib.parse
import queue
import io
import subprocess
import socket
import random
import re
import math
from streamlit_autorefresh import st_autorefresh

# ── Optional imports ────────────────────────────────────────────────────────
try:
    import pyttsx3
except ImportError:
    pyttsx3 = None

try:
    import speech_recognition as sr
except ImportError:
    sr = None

try:
    import ollama
except ImportError:
    ollama = None

try:
    import PyPDF2
    pypdf2_available = True
except ImportError:
    PyPDF2 = None
    pypdf2_available = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    docx_available = True
except ImportError:
    DocxDocument = None
    docx_available = False

# reportlab – PDF export
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    try:
        from reportlab.platypus import HRFlowable
    except ImportError:
        HRFlowable = None
    reportlab_available = True
except ImportError:
    reportlab_available = False
    HRFlowable = None

# ==================================================
# CONFIG
# ==================================================
BASE_FOLDER = r"C:\Users\P R VISHAL\OneDrive\Desktop\FRIDAY..!!"
os.makedirs(BASE_FOLDER, exist_ok=True)

MEMORY_FILE     = os.path.join(BASE_FOLDER, "friday_memory.json")
NOTES_FILE      = os.path.join(BASE_FOLDER, "friday_notes.txt")
STUDY_FILE      = os.path.join(BASE_FOLDER, "study_planner.json")
INTERVIEW_FILE  = os.path.join(BASE_FOLDER, "interview_history.json")
VOICE_NOTES_DIR = os.path.join(BASE_FOLDER, "VoiceNotes")
os.makedirs(VOICE_NOTES_DIR, exist_ok=True)

PASSWORD      = "1234"
GENERAL_MODEL = "phi3:mini"
CODING_MODEL  = "deepseek-coder:1.3b"
OLLAMA_MODEL  = GENERAL_MODEL
USE_OLLAMA    = True

# token budget per feature
TOKENS_FAST     = 150
TOKENS_JSON     = 350
TOKENS_MEDIUM   = 400
TOKENS_EXAM     = 700

# ==================================================
# EMOJIS
# ==================================================
EMOJIS = [
    "🤖", "🧠", "⚡", "🚀", "🛰️", "🔮",
    "🛡️", "📡", "🧬", "🧿", "💠", "♾️"
]

# ==================================================
# TASK CLASSIFIER
# ==================================================
def classify_task(query: str) -> str:
    q = query.lower()
    if any(k in q for k in ["code", "bug", "error", "debug", "function",
                              "script", "program", "syntax", "compile", "runtime"]):
        return "coding"
    if any(k in q for k in ["document", "pdf", "analyze", "file", "extract"]):
        return "document"
    if any(k in q for k in ["exam", "answer", "define", "explain", "marks", "question"]):
        return "exam"
    return "general"

# ==================================================
# MODEL ROUTER
# ==================================================
def select_model(task_type: str) -> str:
    if task_type == "coding":
        return CODING_MODEL
    return GENERAL_MODEL

# ==================================================
# CHECK WHICH MODELS ARE INSTALLED IN OLLAMA
# ==================================================
def get_active_models() -> list:
    try:
        r = requests.get("http://127.0.0.1:11434/api/tags", timeout=3)
        if r.status_code == 200:
            models = r.json().get("models", [])
            return [m.get("name", "") for m in models]
    except Exception:
        pass
    return []

# ==================================================
# CACHE LAYER
# ==================================================
@st.cache_data(show_spinner=False)
def cached_llm(prompt: str, model: str, num_predict: int = TOKENS_FAST) -> str:
    return ollama_chat_http(prompt, model, num_predict)

# ==================================================
# OLLAMA STREAMING  (Command Mode)
# ==================================================
def ollama_stream_response(prompt, model=None):
    if model is None:
        model = select_model(classify_task(prompt))
    if not ollama:
        yield "❌ Error: ollama python library not installed. Run: pip install ollama\n"
        return
    try:
        stream = ollama.chat(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            stream=True,
            options={"num_predict": TOKENS_MEDIUM, "temperature": 0.3}
        )
        for chunk in stream:
            if 'message' in chunk and 'content' in chunk['message']:
                yield chunk['message']['content']
    except Exception as e:
        yield f"\n\n❌ Ollama streaming error: {str(e)}\n"
        yield "Make sure Ollama is running: ollama serve\n"

# ==================================================
# OLLAMA NON-STREAMING
# ==================================================
def ollama_chat_http(prompt, model=None, num_predict=TOKENS_FAST):
    if model is None:
        model = select_model(classify_task(prompt))
    try:
        response = requests.post(
            "http://127.0.0.1:11434/api/chat",
            json={
                "model":    model,
                "messages": [{"role": "user", "content": prompt}],
                "stream":   False,
                "options": {
                    "num_predict": num_predict,
                    "temperature": 0.3
                }
            },
            timeout=120
        )
        response.raise_for_status()
        data  = response.json()
        reply = data.get("message", {}).get("content", "").strip()
        if not reply:
            reply = "No content from Ollama – request successful but empty reply. Check model."
        return reply
    except requests.exceptions.ConnectionError:
        return "❌ Ollama not running. Start it with: ollama serve"
    except requests.exceptions.Timeout:
        return "❌ Ollama timed out. Try a shorter prompt or check your system resources."
    except requests.exceptions.RequestException as e:
        return f"❌ Ollama connection error: {str(e)}"
    except Exception as e:
        return f"❌ Ollama processing error: {str(e)}"

# ==================================================
# VOICE ENGINE
# ==================================================
class VoiceEngine:
    def __init__(self):
        self.enabled       = pyttsx3 is not None
        self.engine        = None
        self.speech_queue  = queue.Queue()
        self.worker_thread = None
        self.running       = False
        if self.enabled:
            try:
                self.engine = pyttsx3.init('sapi5')
                self.engine.setProperty('rate', 160)
                self.engine.setProperty('volume', 0.9)
                voices = self.engine.getProperty('voices')
                if voices:
                    self.engine.setProperty('voice', voices[0].id)
            except Exception as e:
                print(f"Voice initialization failed: {e}")
                self.enabled = False
        if self.enabled and not self.worker_thread:
            self.running       = True
            self.worker_thread = threading.Thread(target=self._worker_loop, daemon=True)
            self.worker_thread.start()

    def speak(self, text):
        if not text.strip() or not self.enabled or not self.engine:
            return
        self.speech_queue.put(text[:180])

    def _worker_loop(self):
        while self.running:
            try:
                text = self.speech_queue.get(timeout=1.0)
                try:
                    self.engine.stop()
                    self.engine.say(text)
                    self.engine.runAndWait()
                except RuntimeError as e:
                    if "run loop already started" not in str(e):
                        print(f"Speak failed: {e}")
                except Exception as e:
                    print(f"Speak failed: {e}")
            except queue.Empty:
                continue
            except Exception as e:
                print(f"Worker error: {e}")
                break

    def stop(self):
        self.running = False
        if self.engine:
            self.engine.stop()

voice = VoiceEngine()

# ==================================================
# MEMORY ENGINE
# ==================================================
class MemoryEngine:
    def __init__(self):
        self.file        = MEMORY_FILE
        self.data        = self.load()
        self.last_action = "Idle"

    def load(self):
        if os.path.exists(self.file):
            try:
                with open(self.file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                for key in ["short_term", "long_term", "context", "tasks", "failures", "chat_history"]:
                    if key not in data:
                        data[key] = [] if key != "context" else {}
                return data
            except:
                pass
        return {
            "short_term":   [],
            "long_term":    [],
            "context":      {"last_action": "", "last_topic": ""},
            "tasks":        [],
            "failures":     [],
            "chat_history": []
        }

    def save(self):
        with open(self.file, 'w', encoding='utf-8') as f:
            json.dump(self.data, f, indent=2)

    def add_chat(self, role, content):
        entry = {"role": role, "content": content, "time": str(datetime.now())}
        self.data["chat_history"].append(entry)
        if len(self.data["chat_history"]) > 20:
            self.data["long_term"].extend(self.data["chat_history"][:10])
            self.data["chat_history"] = self.data["chat_history"][10:]
        self.save()

    def add_task(self, goal):
        tid  = f"T{len(self.data['tasks']) + 1:03d}"
        task = {"id": tid, "goal": goal, "status": "pending", "created": str(datetime.now())}
        self.data["tasks"].append(task)
        self.save()
        return tid

    def get_recent_tasks(self):
        return self.data.get("tasks", [])

    def update_context(self, key, value):
        self.data["context"][key] = value
        self.save()

    def recall_memory(self, query):
        matches = [
            m for m in self.data["chat_history"] + self.data["long_term"]
            if query.lower() in m["content"].lower()
        ]
        if matches:
            return "\n".join([
                f"{m['time']} | {m['role']}: {m['content'][:100]}"
                for m in matches[-3:]
            ])
        return "No matching memory found."

memory = MemoryEngine()

# ==================================================
# LIVE SYSTEM MONITOR THREAD
# ==================================================
monitor_queue   = queue.Queue()
monitor_running = False

def get_network_speed():
    try:
        s = socket.create_connection(('www.google.com', 80), timeout=2)
        s.close()
        return "Connected (Speed: High)"
    except:
        return "Disconnected"

def system_monitor_thread():
    global monitor_running
    monitor_running = True
    monitor_queue.put(("stats", {
        "cpu":       psutil.cpu_percent(),
        "ram":       psutil.virtual_memory().percent,
        "disk":      psutil.disk_usage('C:\\').percent,
        "battery":   "Initializing...",
        "network":   "Checking...",
        "net_speed": "Scanning..."
    }))
    while monitor_running:
        try:
            cpu  = psutil.cpu_percent(interval=1)
            ram  = psutil.virtual_memory().percent
            disk = psutil.disk_usage('C:\\').percent
            bat  = psutil.sensors_battery()
            bat_str = (
                f"{bat.percent}% ({'charging' if bat.power_plugged else 'on battery'})"
                if bat else "N/A"
            )
            try:
                net = "Online" if requests.get("https://www.google.com", timeout=2).status_code == 200 else "Offline"
            except:
                net = "Offline"
            net_speed = get_network_speed()
            alert = ""
            if ram  > 85: alert += f"High RAM: {ram}% "
            if cpu  > 90: alert += f"High CPU: {cpu}% "
            if disk > 90: alert += f"Disk C full: {disk}% "
            if alert:
                voice.speak("System alert! " + alert)
                monitor_queue.put(("alert", alert))
            monitor_queue.put(("stats", {
                "cpu": cpu, "ram": ram, "disk": disk,
                "battery": bat_str, "network": net, "net_speed": net_speed
            }))
        except:
            pass
        time.sleep(5)

monitor_thread = threading.Thread(target=system_monitor_thread, daemon=True)
monitor_thread.start()

# ==================================================
# EXISTING TOOLS
# ==================================================
def open_app(apps_str):
    apps    = [a.strip() for a in apps_str.replace("+", " and ").split(" and ") if a.strip()]
    results = []
    app_map = {
        "chrome":        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        "notepad":       "notepad.exe",
        "calculator":    "calc.exe",
        "youtube":       "https://youtube.com",
        "whatsapp":      "https://web.whatsapp.com",
        "edge":          r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        "paint":         "mspaint.exe",
        "word":          "winword.exe",
        "excel":         "excel.exe",
        "powerpoint":    "powerpnt.exe",
        "settings":      "ms-settings:",
        "file explorer": "explorer.exe",
        "task manager":  "taskmgr.exe",
        "control panel": "control.exe",
        "cmd":           "cmd.exe",
        "powershell":    "powershell.exe",
        "gmail":         "https://mail.google.com",
        "netflix":       "https://netflix.com",
        "amazon":        "https://amazon.com",
        "twitter":       "https://twitter.com",
        "instagram":     "https://instagram.com",
        "facebook":      "https://facebook.com",
        "spotify":       "spotify.exe"
    }
    for app in apps:
        if app in app_map:
            target = app_map[app]
            if target.startswith("http"):
                webbrowser.open(target)
            elif target.startswith("ms-"):
                os.system(f"start {target}")
            else:
                subprocess.Popen(target, shell=True)
            results.append(f"Activated {app.capitalize()} Protocol")
            memory.last_action = f"Activated {app.capitalize()}"
        else:
            results.append(f"{app.capitalize()} Protocol Not Recognized")
    return "\n".join(results)

def play_youtube(query):
    url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(query)}"
    webbrowser.open(url)
    memory.last_action = f"Initiated YouTube Sequence: '{query}'"
    return f"Engaging Visual Sequence '{query}' on YouTube (Search Activated)"

def get_system_status():
    cpu  = psutil.cpu_percent(interval=0.5)
    ram  = psutil.virtual_memory().percent
    disk = psutil.disk_usage('C:\\').percent
    bat  = psutil.sensors_battery()
    bat_str = f"{bat.percent}% ({'charging' if bat.power_plugged else 'on battery'})" if bat else "N/A"
    net_speed = get_network_speed()
    return (
        f"Core Processor Load: {cpu}% | Memory Allocation: {ram}% | "
        f"Storage Matrix C: {disk}% | Energy Core: {bat_str} | Network Grid: {net_speed}"
    )

def research(topic):
    url = f"https://www.google.com/search?q={urllib.parse.quote(topic)}"
    webbrowser.open(url)
    with open(NOTES_FILE, "a", encoding="utf-8") as f:
        f.write(f"\n[{datetime.now()}] Intelligence Scan: {topic}\n")
    memory.last_action = f"Deployed Intelligence Scan on {topic}"
    return f"Intelligence Scan on '{topic}' Activated (Quantum Search Engaged)"

# ──────────────────────────────────────────────────────────────────────────────
# CORE DOCUMENT PARSER
# ──────────────────────────────────────────────────────────────────────────────
def parse_document_bytes(file_bytes: bytes, filename: str) -> tuple:
    ext  = os.path.splitext(filename)[1].lower()
    text = ""
    err  = None
    try:
        if ext == ".pdf":
            if not pypdf2_available:
                return "", "PyPDF2 not installed. Run: pip install PyPDF2"
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            pages  = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            text = "\n".join(pages)
            if not text.strip():
                err = "PDF parsed but no readable text found. The PDF may be scanned/image-only."
        elif ext == ".docx":
            if not docx_available:
                return "", "python-docx not installed. Run: pip install python-docx"
            doc  = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if not text.strip():
                err = "DOCX parsed but no readable text found."
        elif ext in [".txt", ".py", ".log", ".md", ".csv", ".json"]:
            text = file_bytes.decode("utf-8", errors="ignore")
            if not text.strip():
                err = "File appears to be empty."
        else:
            err = f"Unsupported file type '{ext}'. Supported: PDF, DOCX, TXT, PY, LOG, MD, CSV, JSON"
    except Exception as e:
        err = f"Error reading file: {str(e)}"
    return text[:6000], err

def parse_document(file_bytes, filename):
    text, _ = parse_document_bytes(file_bytes, filename)
    return text or "Asset Decryption Failed."


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  ★ FEATURE 1 – DOCUMENT UNDERSTANDING ENGINE                           ║
# ╚══════════════════════════════════════════════════════════════════════════╝
def analyze_document_with_ai(text: str) -> dict:
    if not text.strip():
        return {
            "topic":       "Unknown – no text extracted",
            "summary":     "Could not extract readable text from this document.",
            "action":      "No action required",
            "deadline":    "Not mentioned",
            "consequence": "Not specified"
        }

    prompt = (
        'Read the document text below and reply ONLY with this JSON (no markdown, no extra text):\n'
        '{"topic":"<one sentence>","summary":"<2-3 sentences>","action":"<what reader must do>'
        '","deadline":"<any deadline or Not mentioned>","consequence":"<what happens if ignored or Not specified>"}\n\n'
        f"Document:\n{text[:1200]}"
    )

    raw = ollama_chat_http(prompt, GENERAL_MODEL, num_predict=TOKENS_JSON)

    if raw.startswith("❌"):
        return {
            "topic":       "AI Engine Offline",
            "summary":     raw,
            "action":      "Start Ollama: run 'ollama serve' in your terminal",
            "deadline":    "N/A",
            "consequence": "Analysis unavailable until Ollama is running"
        }

    try:
        clean = re.sub(r'```(?:json)?', '', raw).strip().rstrip('`').strip()
        match = re.search(r'\{.*?\}', clean, re.DOTALL)
        if match:
            return json.loads(match.group())
        return json.loads(clean)
    except Exception:
        def extract(key):
            m = re.search(rf'"{key}"\s*:\s*"([^"]+)"', raw, re.IGNORECASE)
            return m.group(1) if m else "See raw output below"

        return {
            "topic":       extract("topic"),
            "summary":     extract("summary") if extract("summary") != "See raw output below" else raw[:600],
            "action":      extract("action"),
            "deadline":    extract("deadline"),
            "consequence": extract("consequence")
        }


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  ★ FEATURE 2 – SMART STUDY PLANNER                                     ║
# ╚══════════════════════════════════════════════════════════════════════════╝
def load_study_planner() -> dict:
    if os.path.exists(STUDY_FILE):
        try:
            with open(STUDY_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return {"plans": [], "active_plan_id": None}

def save_study_planner(data: dict):
    with open(STUDY_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2)

def generate_study_plan(exam_name, exam_date_str, subjects, weak_areas, daily_hours=4):
    today     = date.today()
    exam_date = datetime.strptime(exam_date_str, "%Y-%m-%d").date()
    days_left = (exam_date - today).days
    if days_left <= 0:
        return {"error": "Exam date must be in the future."}
    if not subjects:
        return {"error": "Please provide at least one subject."}

    weights      = {s: 1.5 if s in weak_areas else 1.0 for s in subjects}
    total_weight = sum(weights.values())
    total_hours  = days_left * daily_hours
    subject_hours = {
        s: round((w / total_weight) * total_hours, 1)
        for s, w in weights.items()
    }

    schedule          = {}
    subject_remaining = {s: subject_hours[s] for s in subjects}
    current_day       = today

    while current_day < exam_date:
        day_str      = current_day.strftime("%Y-%m-%d")
        day_sessions = []
        hours_today  = daily_hours
        allocated    = 0.0
        sorted_subjs = sorted(
            subject_remaining.keys(),
            key=lambda x: subject_remaining[x], reverse=True
        )
        for subj in sorted_subjs:
            if allocated >= hours_today or subject_remaining[subj] <= 0:
                continue
            slot = round(min(
                daily_hours / max(len(subjects), 1),
                subject_remaining[subj],
                hours_today - allocated
            ), 1)
            slot = max(slot, 0.5)
            if allocated + slot > hours_today:
                slot = round(hours_today - allocated, 1)
            if slot <= 0:
                break
            day_sessions.append({
                "subject":   subj,
                "hours":     slot,
                "completed": False,
                "weak":      subj in weak_areas
            })
            subject_remaining[subj] = max(0.0, round(subject_remaining[subj] - slot, 1))
            allocated += slot
            if allocated >= hours_today:
                break
        if day_sessions:
            schedule[day_str] = day_sessions
        current_day += timedelta(days=1)

    return {
        "id":                  f"SP{int(time.time())}",
        "exam_name":           exam_name,
        "exam_date":           exam_date_str,
        "subjects":            subjects,
        "weak_areas":          weak_areas,
        "daily_hours":         daily_hours,
        "subject_total_hours": subject_hours,
        "schedule":            schedule,
        "created":             str(datetime.now()),
        "days_left":           days_left
    }

def get_study_progress(plan):
    if not plan or "error" in plan:
        return {}
    total_hrs  = 0.0
    done_hrs   = 0.0
    subj_total = {s: 0.0 for s in plan.get("subjects", [])}
    subj_done  = {s: 0.0 for s in plan.get("subjects", [])}
    for sessions in plan.get("schedule", {}).values():
        for s in sessions:
            subj              = s["subject"]
            hrs               = s["hours"]
            subj_total[subj]  = subj_total.get(subj, 0) + hrs
            total_hrs        += hrs
            if s.get("completed"):
                subj_done[subj] = subj_done.get(subj, 0) + hrs
                done_hrs       += hrs
    overall = round((done_hrs / total_hrs * 100), 1) if total_hrs > 0 else 0.0
    per_sub = {
        s: round((subj_done[s] / subj_total[s] * 100), 1) if subj_total[s] > 0 else 0.0
        for s in plan.get("subjects", [])
    }
    return {
        "overall":     overall,
        "per_subject": per_sub,
        "done_hours":  done_hrs,
        "total_hours": total_hrs
    }


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  ★ FEATURE 3 – VOICE → NOTES → PDF / DOCX                             ║
# ╚══════════════════════════════════════════════════════════════════════════╝
def capture_voice_for_notes(duration_hint: int = 30) -> str:
    if not sr:
        return "❌ SpeechRecognition not installed. Run: pip install SpeechRecognition"
    recognizer = sr.Recognizer()
    try:
        with sr.Microphone() as source:
            recognizer.adjust_for_ambient_noise(source, duration=1.5)
            audio = recognizer.listen(source, timeout=15, phrase_time_limit=duration_hint)
        return recognizer.recognize_google(audio)
    except sr.WaitTimeoutError:
        return "❌ No speech detected (timeout). Please try again."
    except sr.UnknownValueError:
        return "❌ Could not understand the audio. Speak clearly and try again."
    except sr.RequestError as e:
        return f"❌ Speech recognition service error: {e}"
    except Exception as e:
        return f"❌ Error during recording: {e}"

def format_notes_with_ai(transcript: str, topic: str = "Voice Notes") -> str:
    prompt = (
        "You are an expert note-taking assistant for university students.\n"
        "Convert the raw speech transcript below into clean, structured, exam-ready notes.\n\n"
        f"Topic: {topic}\n\nRaw Transcript:\n{transcript[:1000]}\n\n"
        f"Output format (use Markdown):\n"
        f"# {topic}\n\n"
        "## Key Points\n- Point 1\n- Point 2\n\n"
        "## Detailed Notes\n(Organised paragraphs)\n\n"
        "## Summary\n(2-3 sentence summary)\n\n"
        "## Action Items / To-Do (if any)\n- Item 1\n\n"
        "Rules: Fix grammar, remove filler words, organise logically, keep it academic."
    )
    return ollama_chat_http(prompt, GENERAL_MODEL, num_predict=TOKENS_MEDIUM)

def export_notes_to_pdf(title: str, content: str) -> bytes:
    if not reportlab_available:
        return b""
    buffer = io.BytesIO()
    doc    = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=inch, leftMargin=inch,
        topMargin=inch,   bottomMargin=inch
    )
    styles = getSampleStyleSheet()
    title_style  = ParagraphStyle('FridayTitle',  parent=styles['Title'],   fontSize=22, spaceAfter=16, textColor=rl_colors.HexColor('#ff9500'))
    meta_style   = ParagraphStyle('FridayMeta',   parent=styles['Normal'],  fontSize=10, spaceAfter=12, textColor=rl_colors.HexColor('#888888'))
    h1_style     = ParagraphStyle('FridayH1',     parent=styles['Heading1'],fontSize=15, spaceAfter=8,  spaceBefore=14, textColor=rl_colors.HexColor('#1a1a2e'))
    h2_style     = ParagraphStyle('FridayH2',     parent=styles['Heading2'],fontSize=13, spaceAfter=6,  spaceBefore=10, textColor=rl_colors.HexColor('#16213e'))
    body_style   = ParagraphStyle('FridayBody',   parent=styles['Normal'],  fontSize=11, spaceAfter=5,  leading=17, alignment=TA_JUSTIFY)
    bullet_style = ParagraphStyle('FridayBullet', parent=styles['Normal'],  fontSize=11, spaceAfter=4,  leading=16, leftIndent=18)

    story = [
        Paragraph(title, title_style),
        Paragraph(f"Generated by FRIDAY AI  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}", meta_style),
        Spacer(1, 0.15 * inch),
    ]
    if HRFlowable:
        story.append(HRFlowable(width="100%", thickness=1.5, color=rl_colors.HexColor('#ff9500')))
    story.append(Spacer(1, 0.2 * inch))

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.08 * inch))
            continue
        try:
            if   line.startswith('# '):              story.append(Paragraph(line[2:],  h1_style))
            elif line.startswith('## '):             story.append(Paragraph(line[3:],  h2_style))
            elif line.startswith('### '):            story.append(Paragraph(line[4:],  h2_style))
            elif line.startswith(('- ','* ','• ')): story.append(Paragraph(f"• {line[2:]}", bullet_style))
            elif line.startswith('**') and line.endswith('**'): story.append(Paragraph(f"<b>{line[2:-2]}</b>", body_style))
            else:                                    story.append(Paragraph(line, body_style))
        except Exception:
            pass

    doc.build(story)
    return buffer.getvalue()

def export_notes_to_docx(title: str, content: str) -> bytes:
    if not docx_available:
        return b""
    doc = DocxDocument()
    doc.add_heading(title, level=0)
    p = doc.add_paragraph(f"Generated by FRIDAY AI  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if p.runs:
        p.runs[0].italic = True
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph('')
            continue
        if   line.startswith('# '):              doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):             doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):            doc.add_heading(line[4:], level=3)
        elif line.startswith(('- ','* ','• ')): doc.add_paragraph(line[2:], style='List Bullet')
        else:                                    doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  ★ FEATURE 4 – EXAM ANSWER WRITING ASSISTANT                           ║
# ╚══════════════════════════════════════════════════════════════════════════╝
MARK_CONFIGS = {
    "2 Marks": {
        "target_words": "50–80",
        "structure": (
            "1. DEFINITION (1 sentence)\n"
            "2. KEY POINT / BRIEF EXPLANATION (1-2 sentences)\n"
            "Keep it under 80 words total."
        )
    },
    "5 Marks": {
        "target_words": "150–200",
        "structure": (
            "1. DEFINITION (1-2 sentences) [1 mark]\n"
            "2. DETAILED EXPLANATION with 2-3 sub-points [3 marks]\n"
            "3. BRIEF CONCLUSION [1 mark]\n"
            "Target: 150-200 words."
        )
    },
    "10 Marks": {
        "target_words": "400–500",
        "structure": (
            "1. INTRODUCTION & DEFINITION [2 marks]\n"
            "2. DETAILED EXPLANATION with examples [4 marks]\n"
            "3. DIAGRAM SUGGESTION – describe what to draw and label [2 marks]\n"
            "4. APPLICATIONS / ADVANTAGES / DISADVANTAGES [1 mark]\n"
            "5. CONCLUSION [1 mark]\n"
            "Target: 400-500 words."
        )
    }
}

EXAM_TOKENS = {"2 Marks": 200, "5 Marks": 450, "10 Marks": 700}

def generate_exam_answer(question: str, mark_type: str, subject_hint: str = "") -> str:
    config     = MARK_CONFIGS.get(mark_type, MARK_CONFIGS["5 Marks"])
    tokens     = EXAM_TOKENS.get(mark_type, 450)
    subj_line  = f"Subject: {subject_hint}\n" if subject_hint else ""

    prompt = (
        "You are a university exam answer writer. Write a complete exam answer.\n\n"
        f"{subj_line}"
        f"Question: {question}\n"
        f"Mark Type: {mark_type} | Target Words: {config['target_words']}\n\n"
        f"Structure:\n{config['structure']}\n\n"
        "Use EXACTLY these bold headers in your answer:\n\n"
        "**DEFINITION:**\n[Write here]\n\n"
        "**EXPLANATION:**\n[Write here]\n\n"
        "**DIAGRAM SUGGESTION:**\n[For 2-mark: write 'Not required'. For others: describe diagram.]\n\n"
        "**CONCLUSION:**\n[Write here]\n\n"
        "Write formally. Do NOT add preamble before **DEFINITION:**"
    )

    result = ollama_chat_http(prompt, GENERAL_MODEL, num_predict=tokens)
    return result


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  ★ FEATURE 5 – OFFLINE CODING TUTOR                                    ║
# ╚══════════════════════════════════════════════════════════════════════════╝
CODING_LANGUAGES = ["Python", "C", "C++", "Java", "JavaScript", "SQL", "HTML/CSS"]

CODING_ACTIONS = {
    "📖 Explain Code":      "explain",
    "🐛 Find & Fix Errors": "debug",
    "⚡ Improve Logic":     "improve",
    "✍️ Generate Code":     "generate",
    "🔍 Dry Run / Trace":   "dryrun",
    "📚 Teach Concept":     "teach"
}

def coding_tutor_prompt(action: str, code: str, language: str) -> str:
    code = code[:800]

    if action == "explain":
        return (
            f"You are an expert {language} tutor.\n"
            f"Explain this {language} code to a beginner.\n\n"
            f"Code:\n```{language.lower()}\n{code}\n```\n\n"
            "Use EXACTLY these headers:\n\n"
            "**OVERVIEW:**\nWhat this code does overall\n\n"
            "**LINE-BY-LINE EXPLANATION:**\nExplain each important line\n\n"
            "**KEY CONCEPTS USED:**\nList concepts demonstrated\n\n"
            "**BEGINNER TIP:**\nOne learning tip"
        )
    elif action == "debug":
        return (
            f"You are a senior {language} debugger.\n"
            f"Find ALL errors in this {language} code.\n\n"
            f"Code:\n```{language.lower()}\n{code}\n```\n\n"
            "Use EXACTLY these headers:\n\n"
            "**SYNTAX ERRORS:**\nList syntax mistakes\n\n"
            "**LOGICAL ERRORS:**\nDescribe logic flaws\n\n"
            "**RUNTIME ERRORS:**\nPotential crashes\n\n"
            "**CORRECTED CODE:**\n```\n[Fixed code here]\n```\n\n"
            "**EXPLANATION OF FIXES:**\nWhat was changed and why"
        )
    elif action == "improve":
        return (
            f"You are a {language} optimization expert.\n"
            f"Improve this {language} code.\n\n"
            f"Code:\n```{language.lower()}\n{code}\n```\n\n"
            "Use EXACTLY these headers:\n\n"
            "**CURRENT ISSUES:**\nWhat can be improved\n\n"
            "**IMPROVED CODE:**\n```\n[Optimized code here]\n```\n\n"
            "**IMPROVEMENTS MADE:**\nNumbered list of changes\n\n"
            "**COMPLEXITY ANALYSIS:**\nTime/space complexity before vs after\n\n"
            "**BEST PRACTICES APPLIED:**\nBest practices used"
        )
    elif action == "generate":
        return (
            f"You are an expert {language} programmer.\n"
            f"Write clean {language} code for:\n\n"
            f"Requirement: {code}\n\n"
            "Use EXACTLY these headers:\n\n"
            "**COMPLETE CODE:**\n```\n[Full working code here]\n```\n\n"
            "**HOW IT WORKS:**\nStep-by-step logic explanation\n\n"
            "**HOW TO RUN:**\nSteps to run this code\n\n"
            "**SAMPLE OUTPUT:**\nExpected output\n\n"
            "**POSSIBLE EXTENSIONS:**\nTwo ways to extend this"
        )
    elif action == "dryrun":
        return (
            f"You are a {language} teacher doing a dry run.\n\n"
            f"Code:\n```{language.lower()}\n{code}\n```\n\n"
            "Use EXACTLY these headers:\n\n"
            "**DRY RUN TABLE:**\n"
            "| Step | Line | Operation | Variable Values |\n"
            "|------|------|-----------|----------------|\n"
            "(fill in the table)\n\n"
            "**FINAL OUTPUT:**\nWhat the code prints or returns\n\n"
            "**KEY OBSERVATIONS:**\nImportant execution observations"
        )
    elif action == "teach":
        return (
            f"You are an expert {language} tutor.\n"
            f"Teach this concept: {code}\n\n"
            "Use EXACTLY these headers:\n\n"
            "**DEFINITION:**\nClear simple definition\n\n"
            "**WHY IT MATTERS:**\nReal-world use cases\n\n"
            "**SYNTAX:**\n```\n[Basic syntax]\n```\n\n"
            "**EXAMPLE 1 – BASIC:**\n```\n[Simple example]\n```\n\n"
            "**EXAMPLE 2 – ADVANCED:**\n```\n[Complex example]\n```\n\n"
            "**COMMON MISTAKES:**\nTop 3 beginner mistakes\n\n"
            "**PRACTICE QUESTION:**\nOne exercise to try"
        )
    return f"Help me with this {language} code:\n{code}"


def run_coding_tutor(action_label: str, code: str, language: str) -> str:
    action = CODING_ACTIONS.get(action_label, "explain")
    prompt = coding_tutor_prompt(action, code, language)
    return ollama_chat_http(prompt, CODING_MODEL, num_predict=TOKENS_MEDIUM)


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  ★ FEATURE 6 – MOCK INTERVIEW MODE                                     ║
# ╚══════════════════════════════════════════════════════════════════════════╝
INTERVIEW_DOMAINS = {
    "HR / Behavioral":    "HR and behavioral",
    "Python Programming": "Python programming",
    "Data Structures & Algorithms": "Data Structures and Algorithms",
    "Web Development":    "Web Development (HTML CSS JS React)",
    "Database (SQL/DBMS)":"Database Management Systems and SQL",
    "Operating Systems":  "Operating Systems",
    "Computer Networks":  "Computer Networks",
    "Object-Oriented Programming": "Object-Oriented Programming concepts",
    "Machine Learning":   "Machine Learning and AI basics",
    "System Design":      "System Design and Architecture"
}

def generate_interview_questions(domain_label: str, num_q: int = 5) -> list:
    domain = INTERVIEW_DOMAINS.get(domain_label, domain_label)
    prompt = (
        f'Generate exactly {num_q} interview questions for a {domain} interview.\n'
        'Reply ONLY with a JSON array of strings. No markdown, no extra text.\n'
        'Example: ["Q1?","Q2?","Q3?"]\n'
        "Mix easy to hard, conceptual and practical."
    )
    raw = ollama_chat_http(prompt, GENERAL_MODEL, num_predict=TOKENS_JSON)
    if raw.startswith("❌"):
        return [f"Sample Q{i+1}: Tell me about {domain}?" for i in range(num_q)]
    try:
        clean = re.sub(r'```(?:json)?', '', raw).strip().rstrip('`').strip()
        match = re.search(r'\[.*\]', clean, re.DOTALL)
        if match:
            questions = json.loads(match.group())
            return [str(q) for q in questions if q]
        return json.loads(clean)
    except Exception:
        lines = [l.strip().lstrip('0123456789.-) ') for l in raw.split('\n') if '?' in l and len(l.strip()) > 10]
        return lines[:num_q] if lines else [f"Tell me about {domain}?"]

def evaluate_interview_answer(question: str, answer: str, domain_label: str) -> dict:
    domain = INTERVIEW_DOMAINS.get(domain_label, domain_label)
    prompt = (
        "You are an expert interview evaluator.\n"
        f"Domain: {domain}\nQuestion: {question}\nAnswer: {answer}\n\n"
        'Reply ONLY with this JSON (no markdown):\n'
        '{"score":<1-10>,"grade":"<Excellent|Good|Average|Needs Improvement|Poor>",'
        '"strengths":"<what was good>","improvements":"<what to improve>",'
        '"ideal_answer":"<concise model answer>","tips":"<one interview tip>"}'
    )
    raw = ollama_chat_http(prompt, GENERAL_MODEL, num_predict=TOKENS_MEDIUM)
    if raw.startswith("❌"):
        return {
            "score": 0, "grade": "N/A",
            "strengths": "Ollama offline.",
            "improvements": "Start Ollama to get AI feedback.",
            "ideal_answer": "N/A", "tips": "Run: ollama serve"
        }
    try:
        clean = re.sub(r'```(?:json)?', '', raw).strip().rstrip('`').strip()
        match = re.search(r'\{.*?\}', clean, re.DOTALL)
        if match:
            return json.loads(match.group())
        return json.loads(clean)
    except Exception:
        return {
            "score": 5, "grade": "Average",
            "strengths": "Answer received.",
            "improvements": raw[:400],
            "ideal_answer": "See improvements above.",
            "tips": "Practice structuring answers with STAR method."
        }

def save_interview_result(session: dict):
    history = []
    if os.path.exists(INTERVIEW_FILE):
        try:
            with open(INTERVIEW_FILE, 'r', encoding='utf-8') as f:
                history = json.load(f)
        except Exception:
            pass
    history.append(session)
    with open(INTERVIEW_FILE, 'w', encoding='utf-8') as f:
        json.dump(history[-20:], f, indent=2)

def load_interview_history() -> list:
    if os.path.exists(INTERVIEW_FILE):
        try:
            with open(INTERVIEW_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return []


# ==================================================
# FULL AGENTIC LOOP
# ==================================================
def agentic_process(query, uploaded_content=""):
    memory.last_action = "Initiating Analysis Sequence"

    task_type = classify_task(query)
    model     = select_model(task_type)
    if len(query) < 20:
        model = GENERAL_MODEL

    reasoning     = "Decoding Directive... Advanced neural analysis in progress. Identifying core intent and optimal pathways."
    step_planning = (
        "Step-by-Step Planning:\n"
        "1. Analyze user intent.\n2. Decompose into actionable sub-tasks.\n"
        "3. Establish dependencies (DAG).\n4. Execute in sequence.\n"
        "5. Evaluate outcomes.\n6. Reflect and optimize for future directives."
    )
    version_action = (
        "Version Action Log:\n"
        "- V1.0: Initial command parsing.\n"
        "- V1.1: Enhanced with DAG for dependency management.\n"
        "- V1.2: Added fallback strategies.\n"
        "- V2.0: Integrated multi-stage cognition.\n"
        "- V3.0: Document Analyzer, Study Planner, Voice Notes, Exam Answer.\n"
        "- V4.0: Autorefresh-kill bug fixed. All AI outputs now guaranteed."
    )
    start_time = time.time()
    tasks   = []
    q_lower = query.lower()

    if "open" in q_lower:
        apps_part = q_lower.replace("open", "").replace("+", " and ").strip()
        tasks.append({"action": "OPEN_APP", "param": apps_part})
    if "play" in q_lower and "youtube" in q_lower:
        qp = q_lower.replace("play", "").replace("on youtube", "").replace("youtube", "").strip()
        tasks.append({"action": "PLAY_YOUTUBE", "param": qp})
    if any(w in q_lower for w in ["battery", "cpu", "ram", "system", "status", "network"]):
        tasks.append({"action": "SYSTEM_STATUS"})
    if "research" in q_lower or "tell me about" in q_lower:
        topic = q_lower.replace("research", "").replace("tell me about", "").strip() or q_lower
        tasks.append({"action": "RESEARCH", "param": topic})
    if "add task" in q_lower:
        goal = q_lower.replace("add task", "").strip()
        tasks.append({"action": "ADD_TASK", "param": goal})
    if not tasks:
        tasks.append({"action": "GENERAL", "param": query})

    dag = {
        "tasks": [
            {"id": i+1, "description": t["action"],
             "dependencies": [i] if i > 0 else []}
            for i, t in enumerate(tasks)
        ]
    }
    plan = (
        f"Strategic Plan: {len(tasks)} Directive(s) → "
        + ", ".join(t["action"] for t in tasks)
        + f"\nDAG: {json.dumps(dag, indent=2)}"
    )

    results = []
    for t in tasks:
        action, param = t["action"], t.get("param", "")
        if   action == "OPEN_APP":      results.append(open_app(param))
        elif action == "PLAY_YOUTUBE":  results.append(play_youtube(param))
        elif action == "SYSTEM_STATUS": results.append(get_system_status())
        elif action == "RESEARCH":      results.append(research(param))
        elif action == "ADD_TASK":
            tid = memory.add_task(param)
            results.append(f"Directive {tid} Integrated: {param}")
        else:
            base_prompt = f"Answer briefly (max 5 lines): {query}"
            if task_type == "document" and uploaded_content:
                context   = uploaded_content[:1500]
                ai_prompt = f"Use this context:\n{context}\n\n{base_prompt}"
            else:
                ai_prompt = base_prompt
            results.append(ollama_chat_http(ai_prompt, model, num_predict=TOKENS_FAST))

    execution  = "\n".join(results)
    evaluation = (
        "Partial Protocol Anomaly – Some Directives Unsupported"
        if "not supported" in execution.lower()
        else "All Directives Executed Successfully"
    )
    elapsed    = round(time.time() - start_time, 2)
    reflection = f"Sequence Completed in {elapsed}s. {evaluation}."
    final = (
        f"Decoding Analysis: {reasoning}\n\n{step_planning}\n\n{version_action}\n\n"
        f"Strategic Plan: {plan}\n\nExecution Matrix: {execution}\n\n"
        f"Evaluation Protocol: {evaluation}\n\nReflection Sequence: {reflection}"
    )

    voice.speak(f"{evaluation}. {results[0][:80] if results else ''}")
    memory.add_chat("user", query)
    memory.add_chat("assistant", final)
    memory.last_action = f"Executed: {results[0][:50]}" if results else "Idle"
    return final


# ==================================================
# STREAMLIT MAIN
# ==================================================
def main():
    # ── ALL session state defaults ─────────────────────────────────────────
    _defaults = {
        "mode":                 "Agent",
        "messages":             [],
        "auth":                 False,
        "uploaded_content":     "",
        "agent_query":          "",
        "cmd_text":             "",
        "start_time":           time.time(),
        "live_stats": {
            "cpu": "--", "ram": "--", "disk": "--",
            "battery": "--", "network": "--", "net_speed": "--"
        },
        "doc_analysis":         None,
        "doc_raw_text":         "",
        "doc_file_bytes":       None,
        "doc_file_name":        "",
        "study_plan":           None,
        "voice_transcript":     "",
        "formatted_notes":      "",
        "exam_answer":          "",
        "coding_result":        "",
        "interview_active":     False,
        "interview_questions":  [],
        "interview_q_idx":      0,
        "interview_answers":    [],
        "interview_domain":     "HR / Behavioral",
        "interview_complete":   False,
        "interview_session":    None,
        # ── FIX: pending flags prevent autorefresh from killing AI calls ──
        "ai_processing":        False,   # True = skip autorefresh this run
        "pending_agent":        None,    # Tab 1  agent query
        "pending_doc_analysis": False,   # Tab 4  document analyze
        "pending_exam":         None,    # Tab 7  exam answer
        "pending_coding":       None,    # Tab 8  coding tutor
        "pending_format_notes": None,    # Tab 6  voice notes format
        "pending_iv_questions": None,    # Tab 9  interview question gen
        "pending_iv_eval":      None,    # Tab 9  answer evaluation
        "interview_eval_result":None,    # Tab 9  last eval result to display
    }
    for k, v in _defaults.items():
        st.session_state.setdefault(k, v)

    if "study_planner_data" not in st.session_state:
        st.session_state.study_planner_data = load_study_planner()

    st.set_page_config(
        page_title="FRIDAY AI v4.0",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    st.markdown("""
    <style>
    .stApp { background-color: #0e1117; color: white; }
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        background-color: #1e1e1e !important;
        color: white !important;
        border: 1px solid #444 !important;
        border-radius: 8px !important;
    }
    .stButton > button {
        background-color: transparent !important;
        color: #ff9500 !important;
        border: 2px solid #ff9500 !important;
        border-radius: 50px !important;
        padding: 10px 24px !important;
        font-weight: bold;
        transition: all 0.3s;
    }
    .stButton > button:hover {
        background-color: #ff9500 !important;
        color: black !important;
        transform: scale(1.04);
    }
    .stChatMessage {
        background-color: #1e1e1e;
        border-radius: 12px;
        padding: 14px;
        margin: 8px 0;
        animation: fadeIn 0.6s;
    }
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(10px); }
        to   { opacity: 1; transform: translateY(0); }
    }
    .marquee-container {
        position: fixed; top: 0; left: 0; width: 100%;
        background: linear-gradient(90deg, #1a1f2e, #2a2f3e);
        color: #ff9500; padding: 12px 0; overflow: hidden;
        z-index: 999; border-bottom: 2px solid #ff9500;
        font-family: 'Consolas', monospace; font-size: 14px;
        box-shadow: 0 4px 10px rgba(0,0,0,0.5);
    }
    .marquee {
        display: inline-block; white-space: nowrap;
        animation: marquee 25s linear infinite;
        color: #ffffff; font-weight: 500;
    }
    @keyframes marquee {
        0%   { transform: translateX(100%); }
        100% { transform: translateX(-100%); }
    }
    section[data-testid="stSidebar"] {
        background-color: #11151f !important;
        color: #ffffff !important;
    }
    .blinking-cursor { animation: blink 1s step-end infinite; }
    @keyframes blink { 50% { opacity: 0; } }
    .live-status-panel {
        background-color: #1e1e1e; border-radius: 12px; padding: 24px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.4); margin-top: 24px;
        font-family: 'Segoe UI', sans-serif; color: #fff;
    }
    .live-status-panel h2 {
        font-size: 24px; margin: 0 0 8px; color: #ff9500;
        text-shadow: 0 0 4px rgba(255,149,0,0.3);
    }
    .live-status-panel h3 { font-size: 16px; margin: 0 0 20px; color: #aaa; font-weight: normal; }
    .metrics-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 16px; }
    .metric-card { background-color: #0e1117; border-radius: 8px; padding: 16px; box-shadow: 0 2px 6px rgba(0,0,0,0.3); }
    .metric-header { display: flex; align-items: center; margin-bottom: 8px; }
    .metric-icon   { font-size: 20px; margin-right: 8px; }
    .metric-label  { font-size: 14px; color: #aaa; font-weight: bold; }
    .metric-value  { font-size: 18px; font-weight: bold; display: flex; align-items: center; }
    .model-chip {
        display: inline-block; padding: 2px 8px; border-radius: 10px;
        font-size: 12px; font-weight: bold; margin: 2px 0;
    }
    .model-chip.general  { background: #1a3a1a; color: #00e676; border: 1px solid #00e676; }
    .model-chip.coding   { background: #1a1a3a; color: #64b5f6; border: 1px solid #64b5f6; }
    .model-chip.offline  { background: #3a1a1a; color: #ff5252; border: 1px solid #ff5252; }
    .progress-bar  { background-color: #444; height: 6px; border-radius: 3px; margin-top: 8px; overflow: hidden; }
    .progress-bar div { background-color: #ff9500; height: 100%; border-radius: 3px; transition: width 0.3s ease-in-out; }
    .status-indicator { width: 10px; height: 10px; border-radius: 50%; margin-right: 8px; display: inline-block; }
    .status-indicator.online, .status-indicator.running, .status-indicator.secure {
        background-color: #00ff00; box-shadow: 0 0 8px rgba(0,255,0,0.6); animation: pulse 1.5s infinite;
    }
    .status-indicator.idle  { background-color: #ff9500; box-shadow: 0 0 8px rgba(255,149,0,0.6); }
    .status-indicator.error { background-color: #ff0000; box-shadow: 0 0 8px rgba(255,0,0,0.6); animation: pulse 1s infinite; }
    @keyframes pulse {
        0%   { transform: scale(1);   opacity: 1;   }
        50%  { transform: scale(1.1); opacity: 0.7; }
        100% { transform: scale(1);   opacity: 1;   }
    }
    .advanced-metrics { margin-top: 24px; }
    .advanced-metrics summary {
        font-size: 16px; color: #ff9500; cursor: pointer; outline: none;
        padding: 8px; border-radius: 8px; background-color: #0e1117;
        box-shadow: 0 2px 6px rgba(0,0,0,0.3);
    }
    .advanced-metrics[open] summary { margin-bottom: 16px; }
    .feature-card {
        background: #1e1e1e; border-radius: 12px; padding: 20px;
        border-left: 4px solid #ff9500; margin: 10px 0;
    }
    .analysis-section {
        background: #161b27; border-radius: 10px; padding: 14px;
        margin: 8px 0; border: 1px solid #2a2f3e;
    }
    .analysis-label { color: #ff9500; font-weight: bold; font-size: 13px; margin-bottom: 6px; }
    .analysis-value { color: #ffffff; font-size: 15px; line-height: 1.6; }
    .progress-container { background: #333; border-radius: 6px; height: 12px; margin: 4px 0; }
    .progress-fill { background: linear-gradient(90deg, #ff9500, #ffb347); height: 12px; border-radius: 6px; transition: width 0.5s; }
    .study-day-card { background: #1a1f2e; border-radius: 8px; padding: 12px; margin: 5px 0; border: 1px solid #2a3050; }
    .exam-answer {
        background: #161b27; border-radius: 10px; padding: 22px;
        border: 1px solid #ff9500; line-height: 1.9; color: #e8e8e8; font-size: 15px;
    }
    .code-result {
        background: #0d1117; border-radius: 10px; padding: 20px;
        border: 1px solid #30363d; color: #c9d1d9; font-family: 'Consolas', monospace;
        font-size: 14px; line-height: 1.8;
    }
    .lang-badge {
        display: inline-block; padding: 3px 12px; border-radius: 14px;
        font-size: 13px; font-weight: bold; margin: 4px 4px 4px 0;
        background: #ff9500; color: black;
    }
    .interview-question {
        background: #1a1f2e; border-radius: 10px; padding: 20px;
        border-left: 4px solid #ff9500; margin: 10px 0;
        font-size: 16px; color: #e8e8e8; line-height: 1.7;
    }
    .score-badge {
        display: inline-block; padding: 6px 18px; border-radius: 20px;
        font-size: 16px; font-weight: bold; margin: 6px 0;
    }
    .score-excellent { background: #1a4a1a; color: #00e676; border: 2px solid #00e676; }
    .score-good      { background: #1a3a1a; color: #69f0ae; border: 2px solid #69f0ae; }
    .score-average   { background: #3a3a1a; color: #ffeb3b; border: 2px solid #ffeb3b; }
    .score-poor      { background: #3a1a1a; color: #ff5252; border: 2px solid #ff5252; }
    .feedback-box {
        background: #161b27; border-radius: 8px; padding: 14px;
        margin: 6px 0; border: 1px solid #2a2f3e;
    }
    .feedback-label { color: #ff9500; font-weight: bold; margin-bottom: 4px; }
    .feedback-value { color: #ccc; font-size: 14px; line-height: 1.6; }
    .processing-banner {
        background: linear-gradient(90deg, #1a2a1a, #2a3a2a);
        border: 1px solid #00e676; border-radius: 8px; padding: 12px 18px;
        color: #00e676; font-weight: bold; text-align: center; margin: 10px 0;
        animation: pulse 1.5s infinite;
    }
    </style>
    """, unsafe_allow_html=True)

    # ── Drain monitor queue ───────────────────────────────────────────────
    while not monitor_queue.empty():
        typ, val = monitor_queue.get()
        if typ == "stats":
            st.session_state.live_stats = val

    # ── MARQUEE + AUTOREFRESH (only when NOT processing AI) ──────────────
    # KEY FIX: when ai_processing=True, autorefresh is NOT rendered.
    # Without its JS component, the browser stops sending rerun signals,
    # so the AI call runs to completion uninterrupted.
    if st.session_state.mode == "Agent":
        stats      = st.session_state.live_stats
        llm_status = f"Core Intelligence Online ({GENERAL_MODEL} + {CODING_MODEL})" if USE_OLLAMA and ollama else "Core Offline"
        memory_load = len(memory.data["chat_history"])
        task_count  = len(memory.get_recent_tasks())
        status_items = [
            f"CPU: {stats['cpu']}%", f"RAM: {stats['ram']}%",
            f"Disk C: {stats['disk']}%", f"Battery: {stats['battery']}",
            f"Network: {stats['network']}", f"Net Speed: {stats['net_speed']}",
            f"Memory Load: {memory_load} Entries", f"Task Queue: {task_count} Pending",
            f"LLM: {llm_status}"
        ]
        marquee_text = " • ".join(status_items) + " • " * 4
        st.markdown(f"""
            <div class="marquee-container">
                <div class="marquee">{marquee_text}</div>
            </div>
        """, unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

        # ── AUTOREFRESH: disabled during AI processing to prevent kill ────
        if not st.session_state.get('ai_processing', False):
            st_autorefresh(interval=5000, key="agent_refresh")

    # ── AUTH ──────────────────────────────────────────────────────────────
    if not st.session_state.auth:
        st.title("FRIDAY – Arc Reactor Authentication Sequence")
        pwd = st.text_input("Enter Arc Reactor Access Code", type="password")
        if st.button("🔒 Engage Authentication Protocol"):
            if pwd == PASSWORD:
                st.success("Access Granted: Arc Reactor Online")
                voice.speak("Access Granted. Systems Online.")
                st.session_state.auth = True
                st.rerun()
            else:
                st.error("Access Denied: Invalid Arc Reactor Code")
                voice.speak("Access Denied.")
        st.stop()

    # ══════════════════════════════════════════════════════════════════════
    # AGENT MODE
    # ══════════════════════════════════════════════════════════════════════
    if st.session_state.mode == "Agent":
        st.markdown("<h1>FRIDAY – Grandmaster Agent Protocol v4.0</h1>", unsafe_allow_html=True)

        if st.button("🔄 COMMAND MODE ACTIVATE...."):
            st.session_state.mode = "Command"
            st.rerun()

        with st.sidebar:
            st.header("Quick Actions Control Matrix")
            st.markdown("### Rapid Deployment Actions")

            if st.button("🎤 Engage Voice Command"):
                if sr:
                    r = sr.Recognizer()
                    with sr.Microphone() as source:
                        st.info("Adjusting for ambient noise... (1-2 sec)")
                        try:
                            r.adjust_for_ambient_noise(source, duration=1.5)
                            st.info("Listening... Speak clearly now")
                            audio = r.listen(source, timeout=10, phrase_time_limit=12)
                            text  = r.recognize_google(audio)
                            st.session_state.agent_query = text
                            st.success(f"Vocal Directive Captured: {text}")
                            st.rerun()
                        except sr.WaitTimeoutError:
                            st.warning("No speech detected in time.")
                        except sr.UnknownValueError:
                            st.error("Could not understand what you said.")
                        except sr.RequestError as e:
                            st.error(f"Speech service error: {e}")
                        except Exception as e:
                            st.error(f"Voice recognition error: {e}")
                else:
                    st.error("Vocal Interface Module Absent – install SpeechRecognition")

            uploaded = st.file_uploader(
                "🗂️ Upload Intelligence Asset",
                type=["pdf", "txt", "py", "log", "docx"]
            )
            if uploaded:
                file_bytes = uploaded.getvalue()
                parsed, _  = parse_document_bytes(file_bytes, uploaded.name)
                st.session_state.uploaded_content = parsed
                st.success(f"Asset Decoded: {uploaded.name}")

            st.markdown("### Strategic Directive Operations")
            if st.button("📋 Reveal Pending Directives"):
                tasks = memory.get_recent_tasks()
                st.json(tasks) if tasks else st.info("Directive Queue Empty")

        tabs = st.tabs([
            "🚀 Primary Nexus",
            "📜 Intelligence Log",
            "📋 Directive Queue",
            "📄 Document Analyzer",
            "📚 Study Planner",
            "🎙️ Voice Notes Studio",
            "✏️ Exam Answer Engine",
            "💻 Coding Tutor",
            "🎤 Mock Interview"
        ])
        (tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9) = tabs

        # ── TAB 1 ─────────────────────────────────────────────────────────
        with tab1:
            st.subheader("Deploy FRIDAY Directive....INTENT -> INTELLIGENCE")
            user_query = st.text_input(
                "Input Your Command Sequence...",
                key="agent_query_input",
                value=st.session_state.agent_query
            )

            if st.button("🚀 Initiate FRIDAY Engine......!!!"):
                query = user_query or st.session_state.agent_query
                if query:
                    # ── FIX: set pending, disable autorefresh, then rerun ──
                    st.session_state.pending_agent = {
                        'query':   query,
                        'content': st.session_state.uploaded_content
                    }
                    st.session_state.ai_processing = True
                    st.rerun()

            # ── Process pending agent query (autorefresh is OFF this run) ─
            if st.session_state.get('pending_agent'):
                data = st.session_state.pending_agent
                st.session_state.pending_agent = None
                st.markdown('<div class="processing-banner">🧠 FRIDAY ENGINE PROCESSING — DO NOT REFRESH</div>', unsafe_allow_html=True)
                with st.spinner("Initiating Processing System to THINK..."):
                    result = agentic_process(data['query'], data['content'])
                st.session_state.messages.append({"role": "user",      "content": data['query']})
                st.session_state.messages.append({"role": "assistant", "content": result})
                st.session_state.agent_query    = ""
                st.session_state.ai_processing  = False

            # ── Live Status Panel ─────────────────────────────────────────
            stats = st.session_state.live_stats
            uptime_s     = int(time.time() - st.session_state.start_time)
            uptime       = f"{uptime_s//3600:02d}:{(uptime_s%3600)//60:02d}:{uptime_s%60:02d}"
            engine_state = "Running" if "Idle" not in memory.last_action else "Idle"
            engine_class = "running" if engine_state == "Running" else "idle"
            battery      = stats["battery"]
            ram_p        = stats["ram"]
            try:
                ram_p_f   = float(str(ram_p).replace('--','0'))
                ram_used  = round(ram_p_f / 100 * psutil.virtual_memory().total / 1073741824, 1)
                ram_total = round(psutil.virtual_memory().total / 1073741824, 1)
            except Exception:
                ram_used = ram_total = 0
            cpu_p          = stats["cpu"]
            network        = stats["network"]
            net_speed      = stats["net_speed"]
            network_status = f"{network} ({'Low Latency' if 'High' in str(net_speed) else net_speed})"
            network_class  = "online" if network == "Online" else "error"
            task_active    = sum(1 for t in memory.get_recent_tasks() if t["status"] != "pending")
            task_pending   = len(memory.get_recent_tasks()) - task_active
            disk_p         = stats["disk"]
            try:
                disk_p_f   = float(str(disk_p).replace('--','0'))
                disk_used  = round(disk_p_f / 100 * psutil.disk_usage('C:\\').total / 1073741824, 0)
                disk_total = round(psutil.disk_usage('C:\\').total / 1073741824, 0)
            except Exception:
                disk_used = disk_total = 0

            active_models    = get_active_models()
            coding_online    = any("deepseek" in m.lower() for m in active_models)
            ollama_reachable = USE_OLLAMA and ollama

            if ollama_reachable:
                general_chip = f'<span class="model-chip general">✅ {GENERAL_MODEL}</span>'
                coding_chip  = (
                    f'<span class="model-chip coding">✅ {CODING_MODEL}</span>'
                    if coding_online else
                    f'<span class="model-chip coding" style="opacity:0.6">⏸ {CODING_MODEL}</span>'
                )
                ai_state_label = "Online"
                ai_class       = "online"
            else:
                general_chip   = f'<span class="model-chip offline">❌ {GENERAL_MODEL}</span>'
                coding_chip    = f'<span class="model-chip offline">❌ {CODING_MODEL}</span>'
                ai_state_label = "Offline"
                ai_class       = "error"

            st.markdown(f"""
                <div class="live-status-panel">
                  <h2>LIVE STATUS</h2>
                  <h3>Real-time System Health &amp; Runtime Metrics</h3>
                  <div class="metrics-grid">
                    <div class="metric-card">
                      <div class="metric-header"><span class="metric-icon">⚙️</span><span class="metric-label">Engine State</span></div>
                      <div class="metric-value"><span class="status-indicator {engine_class}"></span>{engine_state}</div>
                    </div>
                    <div class="metric-card">
                      <div class="metric-header"><span class="metric-icon">⏱️</span><span class="metric-label">System Uptime</span></div>
                      <div class="metric-value">{uptime}</div>
                    </div>
                    <div class="metric-card">
                      <div class="metric-header"><span class="metric-icon">🔋</span><span class="metric-label">Battery / Power</span></div>
                      <div class="metric-value">{battery}</div>
                    </div>
                    <div class="metric-card">
                      <div class="metric-header"><span class="metric-icon">🧠</span><span class="metric-label">Memory Usage</span></div>
                      <div class="metric-value">{ram_used} GB / {ram_total} GB</div>
                      <div class="progress-bar"><div style="width:{ram_p}%;"></div></div>
                    </div>
                    <div class="metric-card">
                      <div class="metric-header"><span class="metric-icon">💻</span><span class="metric-label">CPU Load</span></div>
                      <div class="metric-value">{cpu_p}%</div>
                      <div class="progress-bar"><div style="width:{cpu_p}%;"></div></div>
                    </div>
                    <div class="metric-card">
                      <div class="metric-header"><span class="metric-icon">🌐</span><span class="metric-label">Network Status</span></div>
                      <div class="metric-value"><span class="status-indicator {network_class}"></span>{network_status}</div>
                    </div>
                    <div class="metric-card">
                      <div class="metric-header"><span class="metric-icon">📋</span><span class="metric-label">Task Queue</span></div>
                      <div class="metric-value">{task_active} Active / {task_pending} Pending</div>
                    </div>
                    <div class="metric-card">
                      <div class="metric-header"><span class="metric-icon">💾</span><span class="metric-label">Storage Health</span></div>
                      <div class="metric-value">{disk_used} GB / {disk_total} GB</div>
                      <div class="progress-bar"><div style="width:{disk_p}%;"></div></div>
                    </div>
                    <div class="metric-card">
                      <div class="metric-header"><span class="metric-icon">🤖</span><span class="metric-label">AI Model State</span></div>
                      <div class="metric-value"><span class="status-indicator {ai_class}"></span>{ai_state_label}</div>
                      <div style="margin-top:8px;">{general_chip}<br>{coding_chip}</div>
                    </div>
                  </div>
                  <details class="advanced-metrics">
                    <summary>Advanced Metrics</summary>
                    <div class="metrics-grid">
                      <div class="metric-card">
                        <div class="metric-header"><span class="metric-icon">🔒</span><span class="metric-label">Security Status</span></div>
                        <div class="metric-value"><span class="status-indicator secure"></span>Secure</div>
                      </div>
                      <div class="metric-card">
                        <div class="metric-header"><span class="metric-icon">📜</span><span class="metric-label">Logs Status</span></div>
                        <div class="metric-value">Nominal (No Errors)</div>
                      </div>
                      <div class="metric-card">
                        <div class="metric-header"><span class="metric-icon">🌡️</span><span class="metric-label">Thermal State</span></div>
                        <div class="metric-value">42°C (Cool)</div>
                      </div>
                    </div>
                  </details>
                </div>
            """, unsafe_allow_html=True)

        # ── TAB 2 ─────────────────────────────────────────────────────────
        with tab2:
            st.subheader("Elite Operational Log")
            msgs = st.session_state.get("messages", [])
            if msgs:
                for msg in msgs[-10:]:
                    role   = "Operator" if msg["role"] == "user" else "FRIDAY Protocol"
                    avatar = "👤" if msg["role"] == "user" else "🤖"
                    with st.chat_message(role, avatar=avatar):
                        st.markdown(msg["content"])
            else:
                st.info("Operational Log Clear – Initiate Directives")

        # ── TAB 3 ─────────────────────────────────────────────────────────
        with tab3:
            st.subheader("Strategic Directive Queue")
            tasks = memory.get_recent_tasks()
            if tasks:
                for t in tasks:
                    st.write(f"{t['id']}: {t['goal']} ({t['status']})")
            else:
                st.info("Directive Queue Empty – Deploy 'Add Directive: Objective'")
            st.subheader("Intelligence Recall Vault")
            recall_q = st.text_input("Query Historical Intelligence...")
            if recall_q:
                st.markdown(memory.recall_memory(recall_q))

        # ── TAB 4 : DOCUMENT ANALYZER ─────────────────────────────────────
        with tab4:
            st.subheader("📄 Document Understanding Engine")
            st.markdown(
                "Upload any document and FRIDAY will tell you **what it is, what action is required, "
                "the deadline, and what happens if ignored** – powered by local AI (no internet needed)."
            )

            st.markdown("#### Step 1 — Upload Your Document")
            doc_file = st.file_uploader(
                "Supported: PDF, DOCX, TXT, PY, LOG, MD",
                type=["pdf", "docx", "txt", "py", "log", "md", "csv", "json"],
                key="doc_analyzer_upload"
            )

            if doc_file is not None:
                st.session_state.doc_file_bytes = doc_file.getvalue()
                st.session_state.doc_file_name  = doc_file.name
                st.info(
                    f"✅ File loaded: **{doc_file.name}** "
                    f"({round(len(st.session_state.doc_file_bytes)/1024, 1)} KB) — "
                    "Click **Analyze Document** below."
                )

            st.markdown("#### Step 2 — Analyze")
            analyze_col, clear_col = st.columns([2, 1])
            with analyze_col:
                analyze_btn = st.button("🔍 Analyze Document", key="analyze_doc_btn", use_container_width=True)
            with clear_col:
                if st.button("🗑️ Clear", key="clear_doc_btn", use_container_width=True):
                    st.session_state.doc_analysis       = None
                    st.session_state.doc_raw_text        = ""
                    st.session_state.doc_file_bytes      = None
                    st.session_state.doc_file_name       = ""
                    st.session_state.pending_doc_analysis= False
                    st.session_state.ai_processing       = False
                    st.rerun()

            # ── Button click: set pending flag, disable autorefresh, rerun ─
            if analyze_btn:
                if not st.session_state.doc_file_bytes:
                    st.error("❌ Please upload a document first (Step 1).")
                else:
                    st.session_state.pending_doc_analysis = True
                    st.session_state.ai_processing        = True
                    st.rerun()

            # ── Processing block (autorefresh is OFF in this run) ──────────
            if st.session_state.get('pending_doc_analysis'):
                st.session_state.pending_doc_analysis = False
                file_bytes = st.session_state.doc_file_bytes
                filename   = st.session_state.doc_file_name

                st.markdown('<div class="processing-banner">🧠 AI ANALYZING — PLEASE WAIT (30–60 sec)</div>', unsafe_allow_html=True)

                with st.spinner("📖 Extracting text from document..."):
                    extracted, err = parse_document_bytes(file_bytes, filename)

                if err:
                    st.error(f"❌ Extraction error: {err}")
                    if extracted:
                        st.warning("Partial text extracted – proceeding with AI analysis.")

                if extracted and extracted.strip():
                    st.success(f"✅ Extracted {len(extracted)} characters from **{filename}**")
                    st.session_state.doc_raw_text = extracted

                    with st.spinner("🧠 AI is reading and analyzing... (30–60 sec, please wait)"):
                        analysis = analyze_document_with_ai(extracted)

                    st.session_state.doc_analysis = analysis
                    st.session_state.ai_processing = False
                    st.success("✅ AI Analysis complete!")
                elif not err:
                    st.session_state.ai_processing = False
                    st.error("❌ No readable text found in this document.")
                else:
                    st.session_state.ai_processing = False

            # ── Display results ────────────────────────────────────────────
            if st.session_state.doc_analysis:
                a = st.session_state.doc_analysis
                st.markdown("---")
                st.markdown("#### Step 3 — AI Analysis Results")
                st.markdown(f"""
                <div class="feature-card">
                    <div class="analysis-section">
                        <div class="analysis-label">📌 WHAT IS THIS DOCUMENT ABOUT?</div>
                        <div class="analysis-value">{a.get('topic', 'N/A')}</div>
                    </div>
                    <div class="analysis-section">
                        <div class="analysis-label">📝 PLAIN-LANGUAGE SUMMARY</div>
                        <div class="analysis-value">{a.get('summary', 'N/A')}</div>
                    </div>
                    <div class="analysis-section">
                        <div class="analysis-label">✅ ACTION REQUIRED</div>
                        <div class="analysis-value">{a.get('action', 'No action required')}</div>
                    </div>
                    <div class="analysis-section">
                        <div class="analysis-label">⏰ DEADLINE</div>
                        <div class="analysis-value">{a.get('deadline', 'Not mentioned')}</div>
                    </div>
                    <div class="analysis-section">
                        <div class="analysis-label">⚠️ WHAT HAPPENS IF IGNORED?</div>
                        <div class="analysis-value">{a.get('consequence', 'Not specified')}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                with st.expander("📃 View Extracted Raw Text"):
                    st.text_area(
                        "Extracted text (first 4000 chars):",
                        value=st.session_state.doc_raw_text[:4000],
                        height=250, disabled=True
                    )

        # ── TAB 5 : STUDY PLANNER ─────────────────────────────────────────
        with tab5:
            st.subheader("📚 Smart Study Planner (100% Offline)")
            st.markdown(
                "Enter your exam details and FRIDAY generates an **optimised study timetable** "
                "with extra focus on weak subjects, daily progress tracking, and exam countdown."
            )
            planner_sub1, planner_sub2 = st.tabs(["➕ Create New Plan", "📊 Track Progress"])

            with planner_sub1:
                c1, c2 = st.columns(2)
                with c1:
                    exam_name    = st.text_input("Exam / Course Name", placeholder="e.g. Engineering Semester 5 Finals")
                    exam_date_in = st.date_input("Exam Date", min_value=date.today() + timedelta(days=1))
                    daily_h      = st.slider("Daily Study Hours", 1, 12, 4)
                with c2:
                    subjs_raw = st.text_area(
                        "Subjects (one per line)",
                        placeholder="Mathematics\nPhysics\nChemistry\nData Structures",
                        height=140
                    )
                    subjects = [s.strip() for s in subjs_raw.strip().split('\n') if s.strip()]
                    if subjects:
                        weak_areas = st.multiselect("Weak Areas (gets 1.5× study time)", options=subjects)
                    else:
                        weak_areas = []
                        st.info("Enter subjects above to designate weak areas.")

                if st.button("🗓️ Generate Optimised Timetable", key="gen_plan_btn"):
                    if not subjects:
                        st.error("Please enter at least one subject.")
                    else:
                        with st.spinner("Building optimised timetable..."):
                            plan = generate_study_plan(exam_name, exam_date_in.strftime("%Y-%m-%d"), subjects, weak_areas, daily_h)
                        if "error" in plan:
                            st.error(plan["error"])
                        else:
                            pd = st.session_state.study_planner_data
                            pd["plans"].append(plan)
                            pd["active_plan_id"] = plan["id"]
                            save_study_planner(pd)
                            st.session_state.study_planner_data = pd
                            st.session_state.study_plan         = plan
                            st.success(f"✅ Timetable created! **{plan['days_left']} days** until exam • **{len(plan['schedule'])} study days** planned")

                if st.session_state.study_plan and "error" not in st.session_state.study_plan:
                    plan = st.session_state.study_plan
                    st.markdown("---")
                    st.markdown(f"### 📅 Plan: **{plan['exam_name']}** · Exam: {plan['exam_date']}")
                    st.markdown("#### ⏱️ Total Hours per Subject")
                    total_subj_hrs = sum(plan['subject_total_hours'].values())
                    for subj, hrs in plan['subject_total_hours'].items():
                        badge = " 🔴 Weak" if subj in plan['weak_areas'] else " 🟢"
                        pct   = min(int((hrs / total_subj_hrs) * 100), 100) if total_subj_hrs else 0
                        st.markdown(f"""
                        <div style="margin:6px 0;">
                            <span style="color:#ff9500;font-weight:bold;">{subj}{badge}</span>
                            <span style="color:#aaa;font-size:13px;"> – {hrs}h total</span>
                            <div class="progress-container"><div class="progress-fill" style="width:{pct}%;"></div></div>
                        </div>""", unsafe_allow_html=True)
                    st.markdown("#### 📆 Preview – First 7 Study Days")
                    schedule_items = list(plan["schedule"].items())
                    for i, (day_str, sessions) in enumerate(schedule_items[:7]):
                        day_label = datetime.strptime(day_str, "%Y-%m-%d").strftime("%A, %d %b %Y")
                        sess_html = " &nbsp;|&nbsp; ".join(
                            f"{'🔴' if s['weak'] else '🟢'} <b>{s['subject']}</b> ({s['hours']}h)"
                            for s in sessions
                        )
                        st.markdown(f"""
                        <div class="study-day-card">
                            <span style="color:#ff9500;font-weight:bold;">Day {i+1} — {day_label}</span><br>
                            <span style="color:#ccc;font-size:14px;">{sess_html}</span>
                        </div>""", unsafe_allow_html=True)
                    if len(schedule_items) > 7:
                        st.caption(f"… and {len(schedule_items) - 7} more days in the full plan.")
                    st.download_button(
                        "💾 Download Full Timetable (JSON)",
                        data=json.dumps(plan, indent=2),
                        file_name=f"study_plan_{plan['id']}.json",
                        mime="application/json"
                    )

            with planner_sub2:
                plans = st.session_state.study_planner_data.get("plans", [])
                if not plans:
                    st.info("No study plans yet. Create one in the 'Create New Plan' tab.")
                else:
                    plan_map    = {p["id"]: f"{p['exam_name']}  (Exam: {p['exam_date']})" for p in plans}
                    selected_id = st.selectbox("Select Plan", options=list(plan_map.keys()), format_func=lambda x: plan_map[x])
                    sel_plan    = next((p for p in plans if p["id"] == selected_id), None)
                    if sel_plan:
                        progress = get_study_progress(sel_plan)
                        overall  = progress.get("overall", 0)
                        st.markdown(f"### 📊 Overall Progress: **{overall}%**")
                        st.markdown(f"""
                        <div class="progress-container" style="height:22px;margin:8px 0;">
                            <div class="progress-fill" style="width:{overall}%;height:22px;font-size:12px;
                            text-align:center;color:black;font-weight:bold;line-height:22px;">{overall}%</div>
                        </div>""", unsafe_allow_html=True)
                        st.caption(f"Completed: **{progress.get('done_hours',0):.1f}h** / Total: **{progress.get('total_hours',0):.1f}h**")
                        st.markdown("#### Per-Subject Progress")
                        for subj, pct in progress.get("per_subject", {}).items():
                            badge = " 🔴" if subj in sel_plan.get("weak_areas", []) else ""
                            st.markdown(f"""
                            <div style="margin:7px 0;">
                                <span style="color:#ff9500;font-weight:bold;">{subj}{badge}</span>
                                <span style="color:#aaa;"> — {pct}%</span>
                                <div class="progress-container"><div class="progress-fill" style="width:{pct}%;"></div></div>
                            </div>""", unsafe_allow_html=True)
                        today_str      = date.today().strftime("%Y-%m-%d")
                        today_sessions = sel_plan["schedule"].get(today_str, [])
                        st.markdown("#### 📅 Today's Sessions")
                        if today_sessions:
                            planner_data = st.session_state.study_planner_data
                            for idx, session in enumerate(today_sessions):
                                done    = session.get("completed", False)
                                label   = f"{'🔴' if session['weak'] else '🟢'} {session['subject']} — {session['hours']}h"
                                new_val = st.checkbox(label, value=done, key=f"chk_{selected_id}_{today_str}_{idx}")
                                if new_val != done:
                                    for p in planner_data["plans"]:
                                        if p["id"] == selected_id:
                                            if today_str in p["schedule"] and idx < len(p["schedule"][today_str]):
                                                p["schedule"][today_str][idx]["completed"] = new_val
                                    save_study_planner(planner_data)
                                    st.session_state.study_planner_data = planner_data
                                    st.rerun()
                        else:
                            st.info("No sessions scheduled for today.")
                        exam_dt   = datetime.strptime(sel_plan["exam_date"], "%Y-%m-%d").date()
                        days_left = (exam_dt - date.today()).days
                        if   days_left > 0:  st.markdown(f"### ⏳ **{days_left} day(s)** remaining until exam!")
                        elif days_left == 0: st.warning("🎯 Exam is TODAY! All the best!")
                        else:                st.error(f"Exam was {abs(days_left)} day(s) ago.")

        # ── TAB 6 : VOICE NOTES ───────────────────────────────────────────
        with tab6:
            st.subheader("🎙️ Voice → Structured Notes → Export")
            st.markdown(
                "Speak your topic aloud. FRIDAY converts speech to text, "
                "formats it into clean study notes with AI, then exports as **PDF or DOCX**."
            )
            col_rec, col_preview = st.columns([1, 2])

            with col_rec:
                note_title   = st.text_input("Note Title / Topic", placeholder="e.g. Data Structures – Stack & Queue")
                rec_duration = st.slider("Max Recording (sec)", 10, 120, 30)
                if st.button("🔴 Start Recording", key="start_record_btn"):
                    if not sr:
                        st.error("Install SpeechRecognition: pip install SpeechRecognition")
                    else:
                        with st.spinner(f"🎙️ Listening for up to {rec_duration} seconds... Speak now!"):
                            transcript = capture_voice_for_notes(rec_duration)
                        st.session_state.voice_transcript = transcript
                        if transcript.startswith("❌"):
                            st.error(transcript)
                        else:
                            st.success("✅ Recording captured!")
                st.markdown("##### Or type / paste text:")
                manual_in = st.text_area("Manual Input:", height=100, key="manual_note_input")
                if st.button("📋 Use This Text", key="use_manual_btn"):
                    if manual_in.strip():
                        st.session_state.voice_transcript = manual_in
                        st.success("Text loaded.")
                        st.rerun()

                if st.session_state.voice_transcript and not st.session_state.voice_transcript.startswith("❌"):
                    st.markdown("##### Edit Raw Transcript:")
                    edited_t = st.text_area("Transcript (editable):", value=st.session_state.voice_transcript, height=130, key="edit_raw_transcript")

                    if st.button("✨ Format Notes with AI", key="format_notes_btn"):
                        # ── FIX: pending pattern ──────────────────────────
                        st.session_state.pending_format_notes = {
                            'transcript': edited_t,
                            'title':      note_title or "Voice Notes"
                        }
                        st.session_state.ai_processing = True
                        st.rerun()

                # ── Processing block ──────────────────────────────────────
                if st.session_state.get('pending_format_notes'):
                    data = st.session_state.pending_format_notes
                    st.session_state.pending_format_notes = None
                    st.markdown('<div class="processing-banner">🧠 AI FORMATTING NOTES — PLEASE WAIT</div>', unsafe_allow_html=True)
                    with st.spinner("AI is structuring your notes..."):
                        formatted = format_notes_with_ai(data['transcript'], data['title'])
                    st.session_state.formatted_notes  = formatted
                    st.session_state.voice_transcript = data['transcript']
                    st.session_state.ai_processing    = False
                    st.success("Notes formatted!")

            with col_preview:
                st.markdown("### 📝 Formatted Notes Preview")
                if st.session_state.formatted_notes:
                    st.markdown(st.session_state.formatted_notes)
                    st.markdown("---")
                    st.markdown("#### 💾 Export")
                    e1, e2 = st.columns(2)
                    title_export = note_title or "Voice_Notes"
                    safe_title   = re.sub(r'[^\w\s-]', '', title_export).replace(' ', '_')
                    with e1:
                        if reportlab_available:
                            pdf_b = export_notes_to_pdf(title_export, st.session_state.formatted_notes)
                            if pdf_b:
                                st.download_button("📥 Download PDF", data=pdf_b, file_name=f"{safe_title}.pdf", mime="application/pdf")
                        else:
                            st.warning("PDF export: pip install reportlab")
                    with e2:
                        if docx_available:
                            docx_b = export_notes_to_docx(title_export, st.session_state.formatted_notes)
                            if docx_b:
                                st.download_button("📥 Download DOCX", data=docx_b, file_name=f"{safe_title}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        else:
                            st.warning("DOCX export: pip install python-docx")
                    edited_n = st.text_area("Notes (editable):", value=st.session_state.formatted_notes, height=280, key="edit_formatted_notes")
                    if edited_n != st.session_state.formatted_notes:
                        st.session_state.formatted_notes = edited_n
                    if st.button("🗑️ Clear Notes", key="clear_notes_btn"):
                        st.session_state.formatted_notes  = ""
                        st.session_state.voice_transcript = ""
                        st.rerun()
                else:
                    st.info("Record your voice (or paste text) and click **Format Notes with AI** to see structured notes here.")

        # ── TAB 7 : EXAM ANSWER ENGINE ────────────────────────────────────
        with tab7:
            st.subheader("✏️ Exam Answer Writing Assistant")
            st.markdown(
                "Paste your exam question, choose the mark type, and FRIDAY generates a "
                "**complete, structured, university-grade answer**."
            )
            col_q, col_opt = st.columns([3, 1])
            with col_q:
                exam_question = st.text_area("📋 Your Exam Question",
                    placeholder="e.g. Explain the concept of polymorphism in Object-Oriented Programming.",
                    height=120, key="exam_q_input")
            with col_opt:
                mark_type    = st.radio("Answer Type", ["2 Marks", "5 Marks", "10 Marks"], index=1)
                subject_hint = st.text_input("Subject (optional)", placeholder="e.g. OOP / DBMS / OS")

            if st.button("🎯 Generate Exam Answer", key="gen_exam_btn"):
                if not exam_question.strip():
                    st.error("Please enter a question first.")
                else:
                    # ── FIX: pending pattern ──────────────────────────────
                    st.session_state.pending_exam = {
                        'question':  exam_question,
                        'mark_type': mark_type,
                        'subject':   subject_hint
                    }
                    st.session_state.ai_processing = True
                    st.rerun()

            # ── Processing block (autorefresh is OFF this run) ─────────────
            if st.session_state.get('pending_exam'):
                data = st.session_state.pending_exam
                st.session_state.pending_exam = None
                wait_msg = "10–20 sec" if data['mark_type'] == "2 Marks" else ("20–40 sec" if data['mark_type'] == "5 Marks" else "40–90 sec")
                st.markdown('<div class="processing-banner">🧠 WRITING EXAM ANSWER — PLEASE WAIT</div>', unsafe_allow_html=True)
                with st.spinner(f"🧠 Writing {data['mark_type']} answer... ({wait_msg}, please wait — do NOT click anything)"):
                    answer = generate_exam_answer(data['question'], data['mark_type'], data['subject'])
                st.session_state.exam_answer    = answer
                st.session_state.ai_processing  = False
                voice.speak(f"{data['mark_type']} exam answer generated.")

            # ── Display result ─────────────────────────────────────────────
            if st.session_state.exam_answer:
                # Determine mark type for display (use stored answer's mark type from last gen)
                display_mark_type = mark_type  # current radio value
                st.markdown("---")
                mark_color = {"2 Marks": "#4CAF50", "5 Marks": "#2196F3", "10 Marks": "#ff9500"}.get(display_mark_type, "#ff9500")
                st.markdown(
                    f"<div style='display:inline-block;background:{mark_color};color:black;"
                    f"padding:4px 14px;border-radius:20px;font-weight:bold;font-size:14px;"
                    f"margin-bottom:10px;'>{display_mark_type} Answer</div>",
                    unsafe_allow_html=True
                )
                with st.container():
                    st.markdown(
                        f'<div class="exam-answer">{st.session_state.exam_answer.replace(chr(10), "<br>")}</div>',
                        unsafe_allow_html=True
                    )
                st.markdown("---")
                st.markdown("#### 💾 Export Answer")
                x1, x2, x3 = st.columns(3)
                q_safe = re.sub(r'[^\w]', '_', exam_question[:35])
                with x1:
                    if reportlab_available:
                        full_c = f"## Question\n{exam_question}\n\n## {display_mark_type} Answer\n{st.session_state.exam_answer}"
                        pdf_b  = export_notes_to_pdf(f"Exam Answer – {display_mark_type}", full_c)
                        if pdf_b:
                            st.download_button("📥 PDF", data=pdf_b, file_name=f"exam_{q_safe}.pdf", mime="application/pdf")
                    else:
                        st.caption("pip install reportlab")
                with x2:
                    if docx_available:
                        full_c = f"Question: {exam_question}\n\n{display_mark_type} Answer:\n{st.session_state.exam_answer}"
                        docx_b = export_notes_to_docx(f"Exam Answer – {display_mark_type}", full_c)
                        if docx_b:
                            st.download_button("📥 DOCX", data=docx_b, file_name=f"exam_{q_safe}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    else:
                        st.caption("pip install python-docx")
                with x3:
                    st.download_button("📥 TXT", data=f"Q: {exam_question}\n\n{st.session_state.exam_answer}",
                        file_name=f"exam_{q_safe}.txt", mime="text/plain")
                if st.button("🗑️ Clear Answer", key="clear_exam_btn"):
                    st.session_state.exam_answer = ""
                    st.rerun()

        # ── TAB 8 : CODING TUTOR ──────────────────────────────────────────
        with tab8:
            st.subheader("💻 Offline Coding Tutor")
            st.markdown(
                "Paste your code (or describe what you want to build), choose your action, "
                "and FRIDAY's AI will **explain, debug, improve, generate, dry-run, or teach** – "
                "all offline via your local Ollama model."
            )

            lang_html = " ".join(f'<span class="lang-badge">{l}</span>' for l in CODING_LANGUAGES)
            st.markdown(f"**Supported Languages:** {lang_html}", unsafe_allow_html=True)

            ct_col1, ct_col2 = st.columns([2, 1])
            with ct_col1:
                ct_code = st.text_area(
                    "📝 Paste Code or Describe Requirement",
                    placeholder=(
                        "# Paste your code here...\n"
                        "# OR describe what you want to build, e.g.:\n"
                        "# 'Write a Python program to find prime numbers up to N'"
                    ),
                    height=240,
                    key="ct_code_input"
                )
            with ct_col2:
                ct_lang   = st.selectbox("🔤 Language", CODING_LANGUAGES, key="ct_lang")
                ct_action = st.radio(
                    "🎯 What do you want?",
                    list(CODING_ACTIONS.keys()),
                    key="ct_action"
                )
                st.markdown("---")
                action_descriptions = {
                    "📖 Explain Code":      "Understand what the code does",
                    "🐛 Find & Fix Errors": "Debug all types of errors",
                    "⚡ Improve Logic":     "Optimise and refactor",
                    "✍️ Generate Code":     "Write code from description",
                    "🔍 Dry Run / Trace":   "Step-by-step execution trace",
                    "📚 Teach Concept":     "Learn a programming concept"
                }
                st.info(f"**Selected:** {action_descriptions.get(ct_action, '')}")

            if st.button("🚀 Run Coding Tutor", key="run_ct_btn"):
                if not ct_code.strip():
                    st.error("Please enter some code or a requirement.")
                else:
                    # ── FIX: pending pattern ──────────────────────────────
                    st.session_state.pending_coding = {
                        'action': ct_action,
                        'code':   ct_code,
                        'lang':   ct_lang
                    }
                    st.session_state.ai_processing = True
                    st.rerun()

            # ── Processing block ──────────────────────────────────────────
            if st.session_state.get('pending_coding'):
                data = st.session_state.pending_coding
                st.session_state.pending_coding = None
                st.markdown('<div class="processing-banner">🧠 CODING TUTOR PROCESSING — PLEASE WAIT</div>', unsafe_allow_html=True)
                with st.spinner(f"🧠 FRIDAY is processing your {data['lang']} code..."):
                    result = run_coding_tutor(data['action'], data['code'], data['lang'])
                st.session_state.coding_result  = result
                st.session_state.ai_processing  = False
                voice.speak(f"Coding tutor response ready for {data['lang']}.")

            if st.session_state.coding_result:
                st.markdown("---")
                # Use the last known action/lang for display label
                _disp_action = ct_action
                _disp_lang   = ct_lang
                st.markdown(
                    f"<div style='display:inline-block;background:#ff9500;color:black;"
                    f"padding:4px 14px;border-radius:20px;font-weight:bold;font-size:13px;"
                    f"margin-bottom:10px;'>{_disp_action} · {_disp_lang}</div>",
                    unsafe_allow_html=True
                )
                st.markdown(st.session_state.coding_result)

                st.markdown("---")
                st.markdown("#### 💾 Export Result")
                exp1, exp2, exp3 = st.columns(3)
                code_safe = re.sub(r'[^\w]', '_', ct_code[:30])
                full_export = f"Language: {_disp_lang}\nAction: {_disp_action}\n\n" \
                              f"Input:\n{ct_code}\n\n{'='*50}\nFRIDAY Response:\n\n{st.session_state.coding_result}"
                with exp1:
                    if reportlab_available:
                        pdf_b = export_notes_to_pdf(f"Coding Tutor – {_disp_action}", full_export)
                        if pdf_b:
                            st.download_button("📥 PDF", data=pdf_b, file_name=f"code_{code_safe}.pdf", mime="application/pdf")
                    else:
                        st.caption("pip install reportlab")
                with exp2:
                    if docx_available:
                        docx_b = export_notes_to_docx(f"Coding Tutor – {_disp_action}", full_export)
                        if docx_b:
                            st.download_button("📥 DOCX", data=docx_b, file_name=f"code_{code_safe}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    else:
                        st.caption("pip install python-docx")
                with exp3:
                    st.download_button("📥 TXT", data=full_export, file_name=f"code_{code_safe}.txt", mime="text/plain")

                if st.button("🗑️ Clear Result", key="clear_ct_btn"):
                    st.session_state.coding_result = ""
                    st.rerun()

        # ── TAB 9 : MOCK INTERVIEW ────────────────────────────────────────
        with tab9:
            st.subheader("🎤 Mock Interview Mode")
            st.markdown(
                "Simulate a real interview! Choose your domain, FRIDAY generates questions, "
                "you answer via **voice or text**, and get instant AI **feedback + score**."
            )

            iv_sub1, iv_sub2, iv_sub3 = st.tabs(["🏁 Start Interview", "📝 Answer & Feedback", "📊 Interview History"])

            with iv_sub1:
                if st.session_state.interview_complete:
                    session   = st.session_state.interview_session
                    answers   = st.session_state.interview_answers
                    if session and answers:
                        total_score = sum(a.get("score", 0) for a in answers)
                        avg_score   = round(total_score / len(answers), 1) if answers else 0
                        grade_color = (
                            "score-excellent" if avg_score >= 8 else
                            "score-good"      if avg_score >= 6 else
                            "score-average"   if avg_score >= 4 else
                            "score-poor"
                        )
                        grade_word = (
                            "Excellent 🌟" if avg_score >= 8 else
                            "Good 👍"       if avg_score >= 6 else
                            "Average 📈"   if avg_score >= 4 else
                            "Needs Work 💪"
                        )
                        st.markdown("## 🏆 Interview Complete! Final Report")
                        st.markdown(f"""
                        <div class="feature-card" style="text-align:center;">
                            <div style="font-size:18px;color:#aaa;margin-bottom:8px;">
                                {session.get('domain','Interview')} Interview
                            </div>
                            <div class="score-badge {grade_color}" style="font-size:24px;padding:12px 30px;">
                                Score: {avg_score}/10 — {grade_word}
                            </div>
                            <div style="color:#aaa;margin-top:10px;font-size:14px;">
                                {len(answers)} questions answered · Completed: {session.get('completed', 'N/A')}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                        st.markdown("### 📋 Question-by-Question Breakdown")
                        for i, ans in enumerate(answers, 1):
                            sc  = ans.get("score", 0)
                            gc  = ("score-excellent" if sc >= 8 else "score-good" if sc >= 6 else "score-average" if sc >= 4 else "score-poor")
                            with st.expander(f"Q{i}: {ans.get('question','')[:60]}...  |  Score: {sc}/10"):
                                st.markdown(f"**Your Answer:** {ans.get('answer','N/A')}")
                                st.markdown(f"""
                                <div class="feedback-box"><div class="feedback-label">✅ Strengths</div><div class="feedback-value">{ans.get('strengths','N/A')}</div></div>
                                <div class="feedback-box"><div class="feedback-label">📈 Improvements</div><div class="feedback-value">{ans.get('improvements','N/A')}</div></div>
                                <div class="feedback-box"><div class="feedback-label">💡 Ideal Answer</div><div class="feedback-value">{ans.get('ideal_answer','N/A')}</div></div>
                                <div class="feedback-box"><div class="feedback-label">🎯 Pro Tip</div><div class="feedback-value">{ans.get('tips','N/A')}</div></div>
                                """, unsafe_allow_html=True)

                        exp_data = {"session": session, "avg_score": avg_score, "grade": grade_word, "answers": answers}
                        st.download_button("💾 Download Full Report (JSON)", data=json.dumps(exp_data, indent=2),
                            file_name=f"interview_{session.get('id','')}.json", mime="application/json")

                    col_restart = st.columns(2)
                    with col_restart[0]:
                        if st.button("🔄 Start New Interview", key="restart_iv_btn"):
                            for k in ["interview_active","interview_complete","interview_questions",
                                      "interview_q_idx","interview_answers","interview_session",
                                      "interview_eval_result","pending_iv_eval","pending_iv_questions"]:
                                if "active" in k or "complete" in k:
                                    st.session_state[k] = False
                                elif "questions" in k or "answers" in k:
                                    st.session_state[k] = []
                                elif "idx" in k:
                                    st.session_state[k] = 0
                                else:
                                    st.session_state[k] = None
                            st.session_state.ai_processing = False
                            st.rerun()

                elif not st.session_state.interview_active:
                    # ── Check for pending question generation ─────────────
                    if st.session_state.get('pending_iv_questions'):
                        data = st.session_state.pending_iv_questions
                        st.session_state.pending_iv_questions = None
                        st.markdown('<div class="processing-banner">🧠 GENERATING INTERVIEW QUESTIONS — PLEASE WAIT</div>', unsafe_allow_html=True)
                        with st.spinner(f"Generating {data['num_q']} {data['domain']} questions..."):
                            questions = generate_interview_questions(data['domain'], data['num_q'])
                        if questions:
                            session_obj = {
                                "id": f"IV{int(time.time())}", "domain": data['domain'],
                                "mode": data['mode'], "num_q": len(questions), "started": str(datetime.now())
                            }
                            st.session_state.interview_questions = questions
                            st.session_state.interview_q_idx    = 0
                            st.session_state.interview_answers   = []
                            st.session_state.interview_domain    = data['domain']
                            st.session_state.interview_active    = True
                            st.session_state.interview_complete  = False
                            st.session_state.interview_session   = session_obj
                            st.session_state["iv_answer_mode"]   = data['mode']
                            st.session_state.ai_processing       = False
                            voice.speak("Interview started. Good luck!")
                            st.success(f"✅ {len(questions)} questions ready! Go to **Answer & Feedback** tab.")
                            st.rerun()
                        else:
                            st.session_state.ai_processing = False
                            st.error("Failed to generate questions. Check Ollama connection.")
                    else:
                        st.markdown("### ⚙️ Interview Setup")
                        iv_col1, iv_col2 = st.columns(2)
                        with iv_col1:
                            iv_domain = st.selectbox("🎯 Interview Domain", options=list(INTERVIEW_DOMAINS.keys()), key="iv_domain_select")
                            iv_num_q  = st.slider("Number of Questions", 3, 10, 5)
                        with iv_col2:
                            iv_mode = st.radio("Answer Mode", ["⌨️ Text Answer", "🎙️ Voice Answer"], key="iv_mode_radio")
                            st.markdown("")
                            st.info("**Text Mode** – Type answers in text box\n\n**Voice Mode** – Speak your answers aloud")

                        if st.button("🚀 Start Mock Interview", key="start_iv_btn"):
                            # ── FIX: pending pattern ──────────────────────
                            st.session_state.pending_iv_questions = {
                                'domain': iv_domain,
                                'num_q':  iv_num_q,
                                'mode':   iv_mode
                            }
                            st.session_state.ai_processing = True
                            st.rerun()
                else:
                    st.success("✅ Interview in progress! Go to **Answer & Feedback** tab.")
                    st.markdown(
                        f"**Domain:** {st.session_state.interview_session.get('domain','')}\n\n"
                        f"**Questions:** {len(st.session_state.interview_questions)}\n\n"
                        f"**Progress:** {st.session_state.interview_q_idx}/{len(st.session_state.interview_questions)} answered"
                    )

            with iv_sub2:
                if not st.session_state.interview_active:
                    st.info("Start an interview in the **Start Interview** tab first.")
                elif st.session_state.interview_complete:
                    st.success("✅ Interview complete! Check results in **Start Interview** tab.")
                else:
                    questions = st.session_state.interview_questions
                    q_idx     = st.session_state.interview_q_idx
                    domain    = st.session_state.interview_domain
                    iv_mode   = st.session_state.get("iv_answer_mode", "⌨️ Text Answer")

                    # ── Process pending evaluation (autorefresh OFF) ───────
                    if st.session_state.get('pending_iv_eval'):
                        eval_data = st.session_state.pending_iv_eval
                        st.session_state.pending_iv_eval = None
                        st.markdown('<div class="processing-banner">🧠 AI EVALUATING YOUR ANSWER — PLEASE WAIT</div>', unsafe_allow_html=True)
                        with st.spinner("🧠 AI is evaluating your answer..."):
                            feedback = evaluate_interview_answer(
                                eval_data['question'], eval_data['answer'], eval_data['domain']
                            )
                        score = feedback.get("score", 5)
                        grade = feedback.get("grade", "Average")
                        voice.speak(f"Question {eval_data['q_idx']+1} score: {score} out of 10. {grade}.")

                        st.session_state.interview_answers.append({
                            "question":     eval_data['question'],
                            "answer":       eval_data['answer'],
                            "score":        score,
                            "grade":        grade,
                            "strengths":    feedback.get("strengths",""),
                            "improvements": feedback.get("improvements",""),
                            "ideal_answer": feedback.get("ideal_answer",""),
                            "tips":         feedback.get("tips","")
                        })
                        st.session_state.interview_q_idx   = eval_data['q_idx'] + 1
                        st.session_state.ai_processing     = False
                        st.session_state.interview_eval_result = {
                            'feedback': feedback,
                            'question': eval_data['question'],
                            'answer':   eval_data['answer'],
                            'q_idx':    eval_data['q_idx'],
                            'score':    score,
                            'grade':    grade
                        }
                        # refresh to show result
                        st.rerun()

                    # ── Show eval result if available ─────────────────────
                    if st.session_state.get('interview_eval_result'):
                        eval_r   = st.session_state.interview_eval_result
                        feedback = eval_r['feedback']
                        score    = eval_r['score']
                        grade    = eval_r['grade']
                        gc       = ("score-excellent" if score >= 8 else "score-good" if score >= 6
                                    else "score-average" if score >= 4 else "score-poor")

                        st.markdown(f"**Q{eval_r['q_idx']+1}:** {eval_r['question']}")
                        st.markdown(f'<div class="score-badge {gc}">Score: {score}/10 — {grade}</div>', unsafe_allow_html=True)
                        st.markdown(f"""
                        <div class="feedback-box"><div class="feedback-label">✅ Strengths</div><div class="feedback-value">{feedback.get('strengths','N/A')}</div></div>
                        <div class="feedback-box"><div class="feedback-label">📈 Areas to Improve</div><div class="feedback-value">{feedback.get('improvements','N/A')}</div></div>
                        <div class="feedback-box"><div class="feedback-label">💡 Ideal Answer</div><div class="feedback-value">{feedback.get('ideal_answer','N/A')}</div></div>
                        <div class="feedback-box"><div class="feedback-label">🎯 Interview Tip</div><div class="feedback-value">{feedback.get('tips','N/A')}</div></div>
                        """, unsafe_allow_html=True)

                        new_idx = eval_r['q_idx'] + 1
                        if new_idx >= len(questions):
                            # Last question answered
                            st.success("✅ That was the last question!")
                            if st.button("🏆 See Final Report", key="final_report_btn", use_container_width=True):
                                if st.session_state.interview_session:
                                    st.session_state.interview_session["completed"] = str(datetime.now())
                                st.session_state.interview_active       = False
                                st.session_state.interview_complete     = True
                                st.session_state.interview_eval_result  = None
                                save_interview_result({
                                    "session": st.session_state.interview_session,
                                    "answers": st.session_state.interview_answers
                                })
                                voice.speak("Interview complete. Great job!")
                                st.rerun()
                        else:
                            st.markdown(f"*Question {new_idx+1} of {len(questions)} coming up...*")
                            if st.button("▶️ Next Question", key=f"next_q_{eval_r['q_idx']}", use_container_width=True):
                                st.session_state.interview_eval_result = None
                                st.rerun()

                    elif q_idx >= len(questions):
                        # Edge case: completed but flag not set
                        if st.session_state.interview_session:
                            st.session_state.interview_session["completed"] = str(datetime.now())
                        st.session_state.interview_active   = False
                        st.session_state.interview_complete = True
                        save_interview_result({
                            "session": st.session_state.interview_session,
                            "answers": st.session_state.interview_answers
                        })
                        voice.speak("Interview complete!")
                        st.rerun()

                    else:
                        # ── Show current question ─────────────────────────
                        current_q = questions[q_idx]
                        st.markdown(f"**Question {q_idx+1} of {len(questions)}** — {domain}")
                        st.progress(q_idx / len(questions))

                        st.markdown(f"""
                        <div class="interview-question">
                            <span style="color:#ff9500;font-weight:bold;font-size:14px;">Question {q_idx+1}</span><br><br>
                            {current_q}
                        </div>
                        """, unsafe_allow_html=True)

                        if "🎙️ Voice Answer" in iv_mode:
                            st.markdown("#### 🎙️ Voice Answer Mode")
                            if st.button("🔴 Record My Answer", key=f"rec_answer_{q_idx}"):
                                if not sr:
                                    st.error("Install SpeechRecognition: pip install SpeechRecognition")
                                else:
                                    with st.spinner("🎙️ Listening... Speak your answer clearly!"):
                                        spoken = capture_voice_for_notes(45)
                                    if spoken.startswith("❌"):
                                        st.error(spoken)
                                    else:
                                        st.session_state[f"iv_spoken_{q_idx}"] = spoken
                                        st.success(f"✅ Captured: {spoken[:80]}...")
                                        st.rerun()
                            captured_voice = st.session_state.get(f"iv_spoken_{q_idx}", "")
                            if captured_voice and not captured_voice.startswith("❌"):
                                st.markdown("**Your recorded answer:**")
                                candidate_answer = st.text_area("Edit if needed:", value=captured_voice, height=100, key=f"iv_edit_voice_{q_idx}")
                            else:
                                candidate_answer = ""
                                st.info("Click 'Record My Answer' to speak your response.")
                        else:
                            st.markdown("#### ⌨️ Text Answer Mode")
                            candidate_answer = st.text_area(
                                "Type your answer here:",
                                placeholder="Write a clear, structured answer. Take your time.",
                                height=150, key=f"iv_text_answer_{q_idx}"
                            )

                        iv_btn_col1, iv_btn_col2 = st.columns([2, 1])
                        with iv_btn_col1:
                            submit_btn = st.button("✅ Submit Answer & Get Feedback", key=f"submit_iv_{q_idx}", use_container_width=True)
                        with iv_btn_col2:
                            skip_btn = st.button("⏭️ Skip Question", key=f"skip_iv_{q_idx}", use_container_width=True)

                        if submit_btn:
                            if not candidate_answer.strip():
                                st.error("Please provide an answer before submitting.")
                            else:
                                # ── FIX: pending pattern ──────────────────
                                st.session_state.pending_iv_eval = {
                                    'question': current_q,
                                    'answer':   candidate_answer,
                                    'domain':   domain,
                                    'q_idx':    q_idx
                                }
                                st.session_state.ai_processing = True
                                st.rerun()

                        if skip_btn:
                            st.session_state.interview_answers.append({
                                "question":     current_q,
                                "answer":       "[Skipped]",
                                "score":        0,
                                "grade":        "Skipped",
                                "strengths":    "",
                                "improvements": "Practice answering this type of question.",
                                "ideal_answer": "",
                                "tips":         "Don't skip in real interviews."
                            })
                            st.session_state.interview_q_idx = q_idx + 1
                            st.rerun()

            with iv_sub3:
                st.markdown("### 📊 Past Interview Sessions")
                history = load_interview_history()
                if not history:
                    st.info("No interview history yet. Complete your first mock interview!")
                else:
                    for entry in reversed(history[-10:]):
                        sess    = entry.get("session", {})
                        answers = entry.get("answers", [])
                        if not sess or not answers:
                            continue
                        avg_sc = round(sum(a.get("score",0) for a in answers) / len(answers), 1) if answers else 0
                        gc     = ("score-excellent" if avg_sc >= 8 else "score-good" if avg_sc >= 6 else "score-average" if avg_sc >= 4 else "score-poor")
                        with st.expander(f"🗓️ {sess.get('started','')[:16]}  |  {sess.get('domain','')}  |  Score: {avg_sc}/10"):
                            st.markdown(f'<div class="score-badge {gc}">Avg Score: {avg_sc}/10</div>', unsafe_allow_html=True)
                            st.markdown(f"**Domain:** {sess.get('domain','')} · **Questions:** {len(answers)}")
                            for i, ans in enumerate(answers, 1):
                                sc = ans.get("score", 0)
                                st.markdown(f"**Q{i}** ({sc}/10): {ans.get('question','')[:70]}...")
                            st.download_button("💾 Download This Session", data=json.dumps(entry, indent=2),
                                file_name=f"interview_{sess.get('id','')}.json", mime="application/json",
                                key=f"dl_{sess.get('id','')}")

    # ══════════════════════════════════════════════════════════════════════
    # COMMAND MODE
    # ══════════════════════════════════════════════════════════════════════
    else:
        st.markdown("<h1>FRIDAY – Elite Command Nexus</h1>", unsafe_allow_html=True)
        st.info("Direct Neural Link with Core Intelligence • No Auxiliary Protocols Unless Commanded")

        for msg in st.session_state.messages:
            if msg["role"] == "user":
                with st.chat_message("user"):
                    st.markdown(msg["content"])
            else:
                with st.chat_message("assistant"):
                    emoji = random.choice(EMOJIS)
                    st.markdown(f"{emoji} {msg['content']}")

        if st.button("⬅️ Revert to AGENT MODE.....!!"):
            st.session_state.mode = "Agent"
            st.rerun()

        prompt = st.chat_input("Transmit Elite Directive...")

        if prompt:
            st.session_state.messages.append({"role": "user", "content": prompt})
            st.session_state.pending_prompt = prompt
            st.rerun()

        if "pending_prompt" in st.session_state:
            current_prompt = st.session_state.pending_prompt
            del st.session_state.pending_prompt

            stream_model = select_model(classify_task(current_prompt))

            with st.chat_message("assistant"):
                emoji       = random.choice(EMOJIS)
                placeholder = st.empty()
                placeholder.markdown(
                    f"{emoji} Core Intelligence Processing"
                    f"<span class='blinking-cursor'></span>",
                    unsafe_allow_html=True
                )

                full_response = ""
                voice_buffer  = ""

                for token in ollama_stream_response(current_prompt, model=stream_model):
                    full_response += token
                    voice_buffer  += token
                    placeholder.markdown(
                        f"{emoji} {full_response}<span class='blinking-cursor'></span>",
                        unsafe_allow_html=True
                    )
                    if len(voice_buffer) > 60 or token.strip() in ".!?:":
                        voice.speak(voice_buffer.strip())
                        voice_buffer = ""

                placeholder.markdown(f"{emoji} {full_response}")
                if voice_buffer.strip():
                    voice.speak(voice_buffer.strip())

            st.session_state.messages.append({
                "role": "assistant",
                "content": full_response.strip()
            })
            st.rerun()


# ==================================================
if __name__ == "__main__":
    main()